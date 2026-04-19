# ============================================================
# IMPORTS
# ============================================================

import streamlit as st
import pandas as pd
import json
import os
from datetime import datetime
from statistics import mean

# PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from datetime import datetime
import pytz
import time
from zoneinfo import ZoneInfo
import plotly.express as px
import streamlit as st
import json
from datetime import datetime
from github import Github
import time

# ============================================================

# CONFIGURAÇÃO E INICIALIZAÇÃO ÚNICA

# ============================================================

st.set_page_config(

    page_title="Sistema de Análise de Tarefas",

    page_icon="📊",

    layout="wide",

    initial_sidebar_state="expanded"

)

# 2. TRAVA DE SEGURANÇA (Vem logo em seguida)
# ============================================================
st.error("### 🚧 O FORMULÁRIO ENCONTRA-SE INDISPONÍVEL NO MOMENTO.")
st.stop() 
# ============================================================


# Inicialização centralizada

if "logged_in" not in st.session_state: st.session_state.logged_in = False

if "pagina" not in st.session_state:
    st.session_state["pagina"] = "script2"

if "formularios" not in st.session_state: st.session_state["formularios"] = []

       



# Leitura da URL (Prioridade total para permitir acesso ao formulário)

query_params = st.query_params

if "page" in query_params:

    st.session_state.pagina = query_params["page"]

st.markdown("""
    <style>
    /* Oculta a coluna de índice do data_editor */
    div[data-testid="stDataEditor"] > div > div > div > div:first-child {
        display: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

# ============================================================
# 🛡️ INICIALIZAÇÃO DE VARIÁVEIS (PREVINE ERRO 'NOT DEFINED')
# ============================================================
# Aqui dizemos ao Python que essas variáveis existem, mesmo que vazias.
if "nome_f" not in locals(): nome_f = ""
if "cargo" not in locals(): cargo = ""
if "depto" not in locals(): depto = ""
if "setor" not in locals(): setor = ""
if "chefe" not in locals(): chefe = ""
if "unidade" not in locals(): unidade = ""
if "escolaridade" not in locals(): escolaridade = ""
if "devolver_em" not in locals(): devolver_em = ""
if "cursos" not in locals(): cursos = ""
if "objetivo" not in locals(): objetivo = ""
# Caso seu código ainda procure pelo nome antigo em algum lugar:
nome_digitado = st.session_state.get("usuario_atual", "")
# ============================================================



# --- LISTA DE PERGUNTAS DISC ---
perguntas_disc = [
    "Quando surge um problema inesperado: (A) Age rápido | (B) Comunica a todos | (C) Analisa riscos | (D) Segue processo",
    "Em situações de pressão: (A) Foca no resultado | (B) Mantém o otimismo | (C) Mantém a calma | (D) Busca precisão",
    "Ao receber tarefa difícil: (A) Aceita o desafio | (B) Busca ajuda social | (C) Planeja passos | (D) Estuda as regras",
    "No trabalho em equipe: (A) Lidera o grupo | (B) Motiva os colegas | (C) Apoia os outros | (D) Organiza as tarefas",
    "Em reuniões: (A) Vai direto ao ponto | (B) Interage e brinca | (C) Escuta mais | (D) Anota detalhes",
    "Ao lidar com conflitos: (A) Enfrenta direto | (B) Tenta apaziguar | (C) Evita o confronto | (D) Usa lógica e fatos",
    "Seu ritmo de trabalho: (A) Rápido/Impaciente | (B) Rápido/Entusiasmado | (C) Calmo/Constante | (D) Metódico/Cauteloso",
    "Prefere tarefas: (A) Desafiadoras | (B) Variadas e sociais | (C) Rotineiras e seguras | (D) Técnicas e detalhadas",
    "Seu foco principal: (A) Resultados | (B) Relacionamentos | (C) Estabilidade | (D) Qualidade e Processos",
    "Ao decidir, você é: (A) Decidido e firme | (B) Impulsivo e intuitivo | (C) Cuidadoso e lento | (D) Lógico e analítico",
    "Confia mais em: (A) Sua intuição | (B) Opinião alheia | (C) Experiência passada | (D) Dados e provas",
    "Prefere decisões: (A) Independentes | (B) Em grupo | (C) Consensuais | (D) Baseadas em normas",
    "Estilo de organização: (A) Prático | (B) Criativo/Bagunçado | (C) Tradicional | (D) Muito organizado",
    "Lida melhor com: (A) Mudanças rápidas | (B) Novas ideias | (C) Rotinas claras | (D) Regras rígidas",
    "Prefere trabalhar: (A) Sozinho/Comando | (B) Ambiente festivo | (C) Ambiente tranquilo | (D) Ambiente silencioso",
    "Seu ponto forte: (A) Coragem | (B) Comunicação | (C) Paciência | (D) Organização",
    "Você se considera: (A) Dominante | (B) Influente | (C) Estável | (D) Conforme/Analítico",
    "Se motiva por: (A) Poder/Bônus | (B) Reconhecimento | (C) Segurança/Paz | (D) Conhecimento Técnico",
    "Reação a cobranças: (A) Mais esforço | (B) Desculpas criativas | (C) Ansiedade | (D) Argumentos técnicos",
    "Ambiente ideal: (A) Competitivo | (B) Amigável | (C) Previsível | (D) Disciplinado",
    "Ao lidar com feedback: (A) Aceita e ajusta | (B) Comenta e debate | (C) Analisa e planeja | (D) Segue regras",
    "Como prefere aprender: (A) Fazendo | (B) Interagindo | (C) Observando | (D) Estudando materiais",
    "Gestão de tempo: (A) Prioriza resultados | (B) Mantém relações | (C) Planeja com cuidado | (D) Segue processos",
    "Como se comunica: (A) Direto e objetivo | (B) Amigável e motivador | (C) Calmo e ponderado | (D) Técnico e detalhista"
]


# --- FUNÇÕES DE EXPORTAÇÃO (COLE NO TOPO DO SEU ARQUIVO) ---
from docx import Document
from fpdf import FPDF
import io



def extrair_num(texto):
    """Transforma '10 h' ou '5 min' em apenas o número 10 ou 5."""
    try:
        if isinstance(texto, str):
            # Pega apenas os dígitos do texto
            num = "".join(filter(str.isdigit, texto))
            return int(num) if num else 0
        return int(texto)
    except:
        return 0

def limpar_para_rascunho(*args, **kwargs):
    # O (*args, **kwargs) permite que a função receba QUALQUER coisa 
    # e não reclame mais de "arguments".
    st.rerun()

if "respostas_disc" not in st.session_state:
    st.session_state["respostas_disc"] = {}

# ============================================================
# FUNÇÃO DE SUPORTE PARA AS TABELAS (O QUE ESTAVA FALTANDO)
# ============================================================
def preparar_df(chave_rascunho, colunas, fonte_dict):
    """
    Esta função verifica se existe um rascunho carregado.
    Se não, usa os dados oficiais. Se não, cria linhas vazias.
    """
    # 1. Tenta pegar do rascunho carregado no session_state (f_alta_v2, etc)
    # Mapeamento das chaves de rascunho para os nomes das tabelas
    mapa_chaves = {
        "atividades_alta": "f_alta_v2",
        "atividades_normal": "f_normal_v2",
        "atividades_baixa": "f_baixa_v2",
        "dificuldades": "f_dif_v2",
        "sugestoes": "f_sug_v2"
    }
    
    chave_sessao = mapa_chaves.get(chave_rascunho)
    dados_v2 = st.session_state.get(chave_sessao)

    if dados_v2 is not None and len(dados_v2) > 0:
        return pd.DataFrame(dados_v2)

    # 2. Se não tem rascunho, tenta nos dados oficiais (fonte)
    dados_fonte = fonte_dict.get(chave_rascunho, [])
    if dados_fonte:
        # Converte lista simples em DataFrame se necessário
        if isinstance(dados_fonte[0], str):
            return pd.DataFrame([{colunas[0]: item} for item in dados_fonte])
        return pd.DataFrame(dados_fonte)

    # 3. Fallback: Retorna 3 linhas vazias
    return pd.DataFrame([{colunas[0]: ""} for _ in range(3)])

# Configurações de colunas para os editores
col_atv = ["Tarefa"]
col_dif = ["Dificuldade"]
col_sug = ["Sugestão"]
config_col = {"Tarefa": st.column_config.TextColumn("Descrição", width="large")}



@st.cache_data(ttl=10) # 👈 FORÇA O APP A LER O GITHUB A CADA 10 SEGUNDOS
def atualizar_rascunhos_do_github():
    import requests
    import json

    GITHUB_USER = "lucianohcl"
    GITHUB_REPO = "formulario-colaborador"
    FOLDER_PATH = "rascunhos" 
    GITHUB_TOKEN = st.secrets["DB_TOKEN"]

    url = f"https://api.github.com/repos/{GITHUB_USER}/{GITHUB_REPO}/contents/{FOLDER_PATH}"
    headers = {
        "Authorization": f"token {GITHUB_TOKEN}",
        "Accept": "application/vnd.github.v3+json"
    }

    try:
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            arquivos = response.json()
            rascunhos_temp = {}
            
            for arquivo in arquivos:
                if arquivo["name"].endswith(".json"):
                    try:
                        # Adicionamos um parâmetro aleatório na URL para evitar cache do navegador
                        download_url = arquivo["download_url"]
                        conteudo_res = requests.get(download_url, headers=headers)
                        
                        if conteudo_res.status_code != 200:
                            continue
                        
                        dados = conteudo_res.json()
                        if not isinstance(dados, dict):
                            continue

                        # Lógica de resgate do nome (idêntica à sua)
                        nome_colaborador = dados.get("colaborador")
                        if isinstance(nome_colaborador, dict):
                            nome_colaborador = nome_colaborador.get("nome")
                        if not nome_colaborador:
                            nome_colaborador = dados.get("campos", {}).get("nome")
                        
                        if isinstance(nome_colaborador, str):
                            nome_colaborador = nome_colaborador.strip().upper()
                            rascunhos_temp[nome_colaborador] = dados
                    
                    except Exception:
                        continue
            
            # 🔥 ATUALIZAÇÃO SEGURA: Só limpa se realmente trouxe algo novo
            if rascunhos_temp:
                st.session_state["rascunhos"] = rascunhos_temp
                return True
        return False

    except Exception as e:
        # Se der erro de rede, não limpa o que já está na memória
        return False

# --- INICIALIZAÇÃO AUTOMÁTICA ---
if "rascunhos" not in st.session_state:
    st.session_state["rascunhos"] = {}
    atualizar_rascunhos_do_github()


def gerar_word(form):
    doc = Document()
    doc.add_heading(f"Relatório: {form.get('Nome', 'Colaborador')}", 0)
    doc.add_paragraph(f"Data de Envio: {form.get('DataEnvio', 'N/A')}")
    
    # 1. Informações Gerais
    doc.add_heading("Informações de Identificação", level=1)
    campos_gerais = ['Setor', 'Departamento', 'Cargo', 'Chefe', 'Empresa', 'Escolaridade', 'Cursos', 'Objetivo']
    for campo in campos_gerais:
        doc.add_paragraph(f"{campo}: {form.get(campo, 'N/A')}")
    
    # 2. Tabelas (Atividades, Dificuldades, Sugestões)
    secoes = {
        "Atividades": ["Atividade Descrita", "Frequência", "Tempo Gasto"],
        "Dificuldades": ["Dificuldade", "Setor/Parceiro Envolvido", "Tempo Perdido"],
        "Sugestoes": ["Sugestão de Melhoria", "Impacto Esperado"]
    }
    
    for chave, colunas in secoes.items():
        if chave in form and isinstance(form[chave], list):
            doc.add_heading(f"📋 {chave}", level=1)
            # Filtra apenas itens que tenham conteúdo real
            dados = [item for item in form[chave] if any(str(item.get(c, '')).strip() for c in colunas)]
            
            if dados:
                table = doc.add_table(rows=1, cols=len(colunas))
                table.style = 'Table Grid'
                # Cabeçalho
                for i, col in enumerate(colunas):
                    table.rows[0].cells[i].text = col
                # Linhas
                for item in dados:
                    row = table.add_row().cells
                    for i, col in enumerate(colunas):
                        row[i].text = str(item.get(col, ''))
            else:
                doc.add_paragraph("Nenhum dado preenchido nesta seção.")

    # 3. Avaliação DISC
    doc.add_heading("📊 Avaliação DISC (Perguntas e Respostas)", level=1)
    for i, pergunta in enumerate(perguntas_disc, 1):
        valor_resposta = form.get(f"Q{i}", "Não respondido")
        doc.add_paragraph(f"{i}. {pergunta}", style='Heading 2')
        doc.add_paragraph(f"Resposta: {valor_resposta}")
        doc.add_paragraph("-" * 20)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

def gerar_pdf(form):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elementos = []

    # Título
    elementos.append(Paragraph(f"Relatório: {form.get('Nome', 'Colaborador')}", styles['Title']))
    elementos.append(Paragraph(f"Data: {form.get('DataEnvio', 'N/A')}", styles['Normal']))
    elementos.append(Spacer(1, 12))

    # Informações Gerais
    elementos.append(Paragraph("Informações Gerais", styles['Heading2']))
    campos_gerais = ['Setor', 'Departamento', 'Cargo', 'Chefe', 'Empresa', 'Escolaridade', 'Cursos', 'Objetivo']
    for campo in campos_gerais:
        elementos.append(Paragraph(f"<b>{campo}:</b> {form.get(campo, 'N/A')}", styles['Normal']))
    
    elementos.append(Spacer(1, 12))

    # Tabelas (Atividades, Dificuldades, Sugestoes)
    secoes = {
        "Atividades": ["Atividade Descrita", "Frequência", "Tempo Gasto"],
        "Dificuldades": ["Dificuldade", "Setor/Parceiro Envolvido", "Tempo Perdido"],
        "Sugestoes": ["Sugestão de Melhoria", "Impacto Esperado"]
    }
    
    for titulo, colunas in secoes.items():
        if titulo in form and isinstance(form[titulo], list):
            elementos.append(Paragraph(titulo, styles['Heading2']))
            dados = [item for item in form[titulo] if any(str(item.get(c, '')).strip() for c in colunas)]
            
            if dados:
                data = [colunas] # Cabeçalho
                for item in dados:
                    data.append([str(item.get(c, '')) for c in colunas])
                
                tabela = Table(data, repeatRows=1)
                tabela.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                    ('FONTSIZE', (0,0), (-1,-1), 8)
                ]))
                elementos.append(tabela)
            else:
                elementos.append(Paragraph("Nenhum dado preenchido.", styles['Normal']))
            elementos.append(Spacer(1, 12))

    # DISC
    elementos.append(Paragraph("Avaliação DISC", styles['Heading2']))
    for i, pergunta in enumerate(perguntas_disc, 1):
        valor_resposta = form.get(f"Q{i}", "Não respondido")
        elementos.append(Paragraph(f"<b>{i}. {pergunta}</b>", styles['Normal']))
        elementos.append(Paragraph(f"Resposta: {valor_resposta}", styles['Italic']))
        elementos.append(Spacer(1, 6))

import requests

import requests

def enviar_para_sheets(payload):
    url = st.secrets["SHEETS_WEBHOOK"]

    try:
        response = requests.post(url, json=payload)

        st.write("STATUS:", response.status_code)
        st.write("RESPOSTA:", response.text)

        return response.status_code == 200

    except Exception as e:
        st.error(f"Erro ao enviar: {e}")
        return False

# ============================================================
# IMPORTS
# ============================================================

import streamlit as st
import pandas as pd
import json
import os
from datetime import datetime
from statistics import mean

# PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from datetime import datetime
import pytz
import time
from zoneinfo import ZoneInfo
import plotly.express as px
# ============================================================

# CONFIGURAÇÃO E INICIALIZAÇÃO ÚNICA

# ============================================================

st.set_page_config(

    page_title="Sistema de Análise de Tarefas",

    page_icon="📊",

    layout="wide",

    initial_sidebar_state="expanded"

)



# Inicialização centralizada

if "logged_in" not in st.session_state: st.session_state.logged_in = False

if "pagina" not in st.session_state: st.session_state.pagina = "home"

if "formularios" not in st.session_state: st.session_state["formularios"] = []



# Leitura da URL (Prioridade total para permitir acesso ao formulário)

query_params = st.query_params

if "page" in query_params:

    st.session_state.pagina = query_params["page"]

st.markdown("""
    <style>
    /* Oculta a coluna de índice do data_editor */
    div[data-testid="stDataEditor"] > div > div > div > div:first-child {
        display: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

# DEFINE O DIRETÓRIO (Isso resolve o problema da função não achar os arquivos)
dados_dir = "dados"
if not os.path.exists(dados_dir):
    os.makedirs(dados_dir)


# --- LISTA DE PERGUNTAS DISC ---
perguntas_disc = [
    "Quando surge um problema inesperado: (A) Age rápido | (B) Comunica a todos | (C) Analisa riscos | (D) Segue processo",
    "Em situações de pressão: (A) Foca no resultado | (B) Mantém o otimismo | (C) Mantém a calma | (D) Busca precisão",
    "Ao receber tarefa difícil: (A) Aceita o desafio | (B) Busca ajuda social | (C) Planeja passos | (D) Estuda as regras",
    "No trabalho em equipe: (A) Lidera o grupo | (B) Motiva os colegas | (C) Apoia os outros | (D) Organiza as tarefas",
    "Em reuniões: (A) Vai direto ao ponto | (B) Interage e brinca | (C) Escuta mais | (D) Anota detalhes",
    "Ao lidar com conflitos: (A) Enfrenta direto | (B) Tenta apaziguar | (C) Evita o confronto | (D) Usa lógica e fatos",
    "Seu ritmo de trabalho: (A) Rápido/Impaciente | (B) Rápido/Entusiasmado | (C) Calmo/Constante | (D) Metódico/Cauteloso",
    "Prefere tarefas: (A) Desafiadoras | (B) Variadas e sociais | (C) Rotineiras e seguras | (D) Técnicas e detalhadas",
    "Seu foco principal: (A) Resultados | (B) Relacionamentos | (C) Estabilidade | (D) Qualidade e Processos",
    "Ao decidir, você é: (A) Decidido e firme | (B) Impulsivo e intuitivo | (C) Cuidadoso e lento | (D) Lógico e analítico",
    "Confia mais em: (A) Sua intuição | (B) Opinião alheia | (C) Experiência passada | (D) Dados e provas",
    "Prefere decisões: (A) Independentes | (B) Em grupo | (C) Consensuais | (D) Baseadas em normas",
    "Estilo de organização: (A) Prático | (B) Criativo/Bagunçado | (C) Tradicional | (D) Muito organizado",
    "Lida melhor com: (A) Mudanças rápidas | (B) Novas ideias | (C) Rotinas claras | (D) Regras rígidas",
    "Prefere trabalhar: (A) Sozinho/Comando | (B) Ambiente festivo | (C) Ambiente tranquilo | (D) Ambiente silencioso",
    "Seu ponto forte: (A) Coragem | (B) Comunicação | (C) Paciência | (D) Organização",
    "Você se considera: (A) Dominante | (B) Influente | (C) Estável | (D) Conforme/Analítico",
    "Se motiva por: (A) Poder/Bônus | (B) Reconhecimento | (C) Segurança/Paz | (D) Conhecimento Técnico",
    "Reação a cobranças: (A) Mais esforço | (B) Desculpas criativas | (C) Ansiedade | (D) Argumentos técnicos",
    "Ambiente ideal: (A) Competitivo | (B) Amigável | (C) Previsível | (D) Disciplinado",
    "Ao lidar com feedback: (A) Aceita e ajusta | (B) Comenta e debate | (C) Analisa e planeja | (D) Segue regras",
    "Como prefere aprender: (A) Fazendo | (B) Interagindo | (C) Observando | (D) Estudando materiais",
    "Gestão de tempo: (A) Prioriza resultados | (B) Mantém relações | (C) Planeja com cuidado | (D) Segue processos",
    "Como se comunica: (A) Direto e objetivo | (B) Amigável e motivador | (C) Calmo e ponderado | (D) Técnico e detalhista"
]

# --- FUNÇÕES DE EXPORTAÇÃO (COLE NO TOPO DO SEU ARQUIVO) ---
from docx import Document
from fpdf import FPDF
import io

def salvar_no_github(payload, nome_arquivo, pasta="rascunhos"):
    from github import Github
    import json

    # ✅ Usa o token corretamente do st.secrets
    GITHUB_TOKEN = st.secrets["DB_TOKEN"]
    REPO_NOME = "lucianohcl/formulario-colaborador"

    g = Github(GITHUB_TOKEN)
    repo = g.get_repo(REPO_NOME)

    caminho = f"{pasta}/{nome_arquivo}"  # junta a pasta + nome do arquivo
    try:
        conteudo = json.dumps(payload, indent=4, ensure_ascii=False)
        try:
            file = repo.get_contents(caminho)
            repo.update_file(caminho, f"Atualizando {nome_arquivo}", conteudo, file.sha)
        except:
            repo.create_file(caminho, f"Criando {nome_arquivo}", conteudo)
        return True
    except Exception as e:
        st.error(f"❌ Erro ao salvar no GitHub: {e}")
        return False

def gerar_word(form):
    doc = Document()
    doc.add_heading(f"Relatório: {form.get('Nome', 'Colaborador')}", 0)
    doc.add_paragraph(f"Data de Envio: {form.get('DataEnvio', 'N/A')}")
    
    # 1. Informações Gerais
    doc.add_heading("Informações de Identificação", level=1)
    campos_gerais = ['Setor', 'Departamento', 'Cargo', 'Chefe', 'Empresa', 'Escolaridade', 'Cursos', 'Objetivo']
    for campo in campos_gerais:
        doc.add_paragraph(f"{campo}: {form.get(campo, 'N/A')}")
    
    # 2. Tabelas (Atividades, Dificuldades, Sugestões)
    secoes = {
        "Atividades": ["Atividade Descrita", "Frequência", "Tempo Gasto"],
        "Dificuldades": ["Dificuldade", "Setor/Parceiro Envolvido", "Tempo Perdido"],
        "Sugestoes": ["Sugestão de Melhoria", "Impacto Esperado"]
    }
    
    for chave, colunas in secoes.items():
        if chave in form and isinstance(form[chave], list):
            doc.add_heading(f"📋 {chave}", level=1)
            # Filtra apenas itens que tenham conteúdo real
            dados = [item for item in form[chave] if any(str(item.get(c, '')).strip() for c in colunas)]
            
            if dados:
                table = doc.add_table(rows=1, cols=len(colunas))
                table.style = 'Table Grid'
                # Cabeçalho
                for i, col in enumerate(colunas):
                    table.rows[0].cells[i].text = col
                # Linhas
                for item in dados:
                    row = table.add_row().cells
                    for i, col in enumerate(colunas):
                        row[i].text = str(item.get(col, ''))
            else:
                doc.add_paragraph("Nenhum dado preenchido nesta seção.")

    # 3. Avaliação DISC
    doc.add_heading("📊 Avaliação DISC (Perguntas e Respostas)", level=1)
    for i, pergunta in enumerate(perguntas_disc, 1):
        valor_resposta = form.get(f"Q{i}", "Não respondido")
        doc.add_paragraph(f"{i}. {pergunta}", style='Heading 2')
        doc.add_paragraph(f"Resposta: {valor_resposta}")
        doc.add_paragraph("-" * 20)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

def gerar_pdf(form):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elementos = []

    # Título
    elementos.append(Paragraph(f"Relatório: {form.get('Nome', 'Colaborador')}", styles['Title']))
    elementos.append(Paragraph(f"Data: {form.get('DataEnvio', 'N/A')}", styles['Normal']))
    elementos.append(Spacer(1, 12))

    # Informações Gerais
    elementos.append(Paragraph("Informações Gerais", styles['Heading2']))
    campos_gerais = ['Setor', 'Departamento', 'Cargo', 'Chefe', 'Empresa', 'Escolaridade', 'Cursos', 'Objetivo']
    for campo in campos_gerais:
        elementos.append(Paragraph(f"<b>{campo}:</b> {form.get(campo, 'N/A')}", styles['Normal']))
    
    elementos.append(Spacer(1, 12))

    # Tabelas (Atividades, Dificuldades, Sugestoes)
    secoes = {
        "Atividades": ["Atividade Descrita", "Frequência", "Tempo Gasto"],
        "Dificuldades": ["Dificuldade", "Setor/Parceiro Envolvido", "Tempo Perdido"],
        "Sugestoes": ["Sugestão de Melhoria", "Impacto Esperado"]
    }
    
    for titulo, colunas in secoes.items():
        if titulo in form and isinstance(form[titulo], list):
            elementos.append(Paragraph(titulo, styles['Heading2']))
            dados = [item for item in form[titulo] if any(str(item.get(c, '')).strip() for c in colunas)]
            
            if dados:
                data = [colunas] # Cabeçalho
                for item in dados:
                    data.append([str(item.get(c, '')) for c in colunas])
                
                tabela = Table(data, repeatRows=1)
                tabela.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                    ('FONTSIZE', (0,0), (-1,-1), 8)
                ]))
                elementos.append(tabela)
            else:
                elementos.append(Paragraph("Nenhum dado preenchido.", styles['Normal']))
            elementos.append(Spacer(1, 12))

    # DISC
    elementos.append(Paragraph("Avaliação DISC", styles['Heading2']))
    for i, pergunta in enumerate(perguntas_disc, 1):
        valor_resposta = form.get(f"Q{i}", "Não respondido")
        elementos.append(Paragraph(f"<b>{i}. {pergunta}</b>", styles['Normal']))
        elementos.append(Paragraph(f"Resposta: {valor_resposta}", styles['Italic']))
        elementos.append(Spacer(1, 6))

    doc.build(elementos)
    buffer.seek(0)
    return buffer

# ============================================================
# CALCULAR DISC PERCENTUAL E DOMINANTE
# ============================================================

def calcular_disc(respostas_disc):
    contagem = {"D":0, "I":0, "S":0, "C":0}
    for r in respostas_disc.values():
        if r in contagem:
            contagem[r] += 1
    total = sum(contagem.values())
    if total > 0:
        percentuais = {k: round(v/total*100,1) for k,v in contagem.items()}
        dominante = max(percentuais, key=percentuais.get)
    else:
        percentuais = contagem
        dominante = None
    return percentuais, dominante

# ============================================================
# SCORE DISC PONDERADO
# ============================================================

def score_disc(disc):
    pesos = {"D":1.0,"I":0.9,"S":0.85,"C":0.95}
    total = sum(disc.values())
    if total == 0:
        return 0
    calculo = sum(disc[k]*pesos.get(k,1) for k in disc)
    return round((calculo/total)*100,2)


# ============================================================
# DEFINIÇÃO E CARREGAMENTO DO BANCO DE DADOS (AJUSTADO)
# ============================================================
import streamlit as st
import pandas as pd
import os
import json
import sys

import os
import sys
import json
import streamlit as st

# --- DEFINIÇÃO DE CAMINHO À PROVA DE ERROS ---
if getattr(sys, 'frozen', False):
    base_dir = os.path.dirname(sys.executable)
else:
    base_dir = os.path.dirname(os.path.abspath(__file__))

# Definimos o diretório de dados como absoluto
dados_dir = os.path.join(base_dir, "dados")

# Criamos a pasta 'dados' se ela não existir
os.makedirs(dados_dir, exist_ok=True)

# --- FUNÇÃO DE CARREGAMENTO DINÂMICO ---
def carregar_todos_formularios():
    """
    Lê todos os arquivos .json da pasta 'dados' individualmente.
    """
    lista_formularios = []
    # Usamos a variável global dados_dir definida acima
    if os.path.exists(dados_dir):
        for nome_arquivo in os.listdir(dados_dir):
            if nome_arquivo.endswith(".json"):
                caminho_completo = os.path.join(dados_dir, nome_arquivo)
                try:
                    with open(caminho_completo, "r", encoding="utf-8") as f:
                        dados = json.load(f)
                        if isinstance(dados, dict):
                            lista_formularios.append(dados)
                except Exception:
                    continue
    return lista_formularios

# --- CARREGAMENTO INICIAL ---
# Agora chamamos a função que criamos para ler os arquivos individuais
if "formularios" not in st.session_state:
    st.session_state["formularios"] = carregar_todos_formularios()
# ============================================================
# LOGIN (Com Bypass para o Formulário)
# ============================================================
# Só bloqueia o acesso se NÃO estiver logado E NÃO for a página de formulário
if not st.session_state.logged_in and st.session_state.pagina != "formulario":
    st.title("🔐 Acesso")
    usuario = st.text_input("Usuário")
    senha = st.text_input("Senha", type="password")

    if st.button("Entrar", key="login_button"):
        if (usuario == "admin" and senha == "admin123") or (usuario == "Luciano" and senha == "123"):
            st.session_state.logged_in = True
            st.session_state.user_nome = usuario
            st.session_state.is_admin = True
            
            # ATUALIZAÇÃO: Definimos a variável que o painel de exportação espera
            if usuario == "Luciano":
                st.session_state["usuario_logado"] = "Luciano 123"
            else:
                st.session_state["usuario_logado"] = usuario
                
            st.rerun()
        else:
            st.error("Usuário ou senha incorretos")
    
    st.stop()

# ============================================================
# SIDEBAR
# ============================================================

st.sidebar.title("📌 Menu de Navegação")

btn_home = st.sidebar.button("🏠 Home")
btn_analise = st.sidebar.button("📊 Análise Inteligente")
btn_comparar = st.sidebar.button("⚖️ Comparar Colaboradores")
btn_disc = st.sidebar.button("🧠 Perfil DISC")
btn_parecer = st.sidebar.button("📄 Parecer Estratégico")
btn_visualizar = st.sidebar.button("👁️ Visualizar Dados")
btn_produtividade = st.sidebar.button("🚀 Produtividade")


st.sidebar.markdown("---")

btn_logout = st.sidebar.button("🚪 Logout")

pagina_anterior = st.session_state.pagina

if btn_home:
    st.session_state.pagina = "home"
elif btn_analise:
    st.session_state.pagina = "analise"
elif btn_comparar:
    st.session_state.pagina = "comparar"
elif btn_disc:
    st.session_state.pagina = "disc"
elif btn_parecer:
    st.session_state.pagina = "parecer"
elif btn_visualizar:
    st.session_state.pagina = "visualizar"
# O elif abaixo verifica a URL sem precisar de botão
elif st.session_state.pagina == "formulario":
    pass # Este comando é obrigatório para não dar erro de sintaxe
elif btn_logout:
    st.session_state.logged_in = False
    st.session_state.pagina = "home"

if pagina_anterior != st.session_state.pagina:
    st.rerun()

# ============================================================
# PÁGINA PERFIL DISC (VERSÃO SINCRO)
# ============================================================

if st.session_state.pagina == "disc":
    import plotly.express as px
    import pandas as pd

    st.title("🧠 Análise de Perfil DISC")

    # 1. FORÇAR LEITURA DIRETA (IGUAL AO VISUALIZAR REGISTROS)
    # Isso garante que não dependemos de um session_state que pode estar vazio
    lista_fresca = carregar_todos_formularios()

    if not lista_fresca:
        st.warning("Nenhum formulário encontrado na pasta de dados.")
        if st.button("♻️ Tentar recarregar dados"):
            st.rerun()
        st.stop()

    # 2. MAPEAMENTO SEGURO
    opcoes_colaboradores = {
        f"{f.get('nome', 'Sem Nome')} - {f.get('cargo', 'Sem Cargo')}": f 
        for f in lista_fresca
    }

    colaborador_chave = st.selectbox(
        "Escolha o colaborador",
        options=list(opcoes_colaboradores.keys())
    )

    # 3. RECUPERAÇÃO DO FORMULÁRIO
    formulario_sel = opcoes_colaboradores.get(colaborador_chave)

    # ============================================================
    # BOTÃO GERAR ANÁLISE
    # ============================================================

    if formulario_sel and st.button("🔎 Gerar análise DISC"):
        # A partir daqui o seu processamento continua normal
        form = formulario_sel
        
        mapa_disc = {
            "A": "D",
            "B": "I",
            "C": "S",
            "D": "C"
        }
        

        # Extraímos as respostas garantindo que o dicionário 'disc' existe no JSON
        respostas_raw = form.get("disc", {})
        respostas_disc = {}

        for k, v in respostas_raw.items():
            if v in mapa_disc:
                respostas_disc[k] = mapa_disc[v]

        # ============================================================
        # PAINEL DISC DO COLABORADOR (AJUSTADO)
        # ============================================================

        # 1️⃣ Função ajustada de cálculo de score
        def score_disc(percentuais):
            """
            Calcula a intensidade do perfil dominante considerando a diferença
            entre ele e o segundo maior perfil.
            Retorna um valor de 0 a 100, refletindo a certeza relativa.
            """
            if not percentuais:
                return 0
            
            valores = sorted(percentuais.values(), reverse=True)
            dominante_val = valores[0]
            segundo_val = valores[1] if len(valores) > 1 else 0
            
            diff = dominante_val - segundo_val
            score_normalizado = round((diff / dominante_val) * 100, 1) if dominante_val > 0 else 0
            score_normalizado = max(0, min(score_normalizado, 100))
            
            return score_normalizado

        # 2️⃣ Cálculos
        percentuais, dominante = calcular_disc(respostas_disc)
        score = score_disc(percentuais)

        st.markdown("## 🔹 Painel DISC do Colaborador")

        # 3️⃣ Gráfico e Métricas lado a lado
        col_graf, col_met = st.columns([2,1])

        with col_graf:
            fig = px.bar(
                x=list(percentuais.keys()),
                y=list(percentuais.values()),
                labels={'x':'Tipo','y':'Percentual (%)'},
                text=list(percentuais.values()),
                color=list(percentuais.keys()),
                color_discrete_map={"D":"#FF4136","I":"#FF851B","S":"#2ECC40","C":"#0074D9"}
            )
            fig.update_layout(
                yaxis_range=[0,100], 
                height=350, 
                margin=dict(l=20, r=20, t=30, b=20), 
                template="plotly_white",
                showlegend=False
            )
            st.plotly_chart(fig, use_container_width=True)

        with col_met:
            st.metric("Perfil Dominante", dominante)
            st.metric("Intensidade (Score)", f"{score}%")
            
            # Interpretação rápida do nível de intensidade
            def interpretar_valor(p):
                try:
                    v = float(str(p).replace('%',''))
                    if v > 85: return "🎯 **Muito Alta**"
                    if v > 60: return "✅ **Alta**"
                    if v > 30: return "⚖️ **Moderada**"
                    return "⚠️ **Baixa**"
                except:
                    return ""
            
            st.write(interpretar_valor(score))

            st.caption("ℹ️ Score indica a intensidade relativa do perfil dominante em relação aos outros perfis. Quanto maior a diferença, maior a certeza do perfil.")


            st.markdown("---")

            

        # 2. INTERPRETAÇÃO DETALHADA (Substitui a Base de Conhecimento e o Parecer)
        textos_disc = {
            "D": {"nome": "Dominante", "estilo": "Resultados e Assertividade", "desc": "Decidido e direto. Busca desafios e rapidez.", "cor": "red", "tarefas": "Tomada de decisão, Gestão de crises, Metas."},
            "I": {"nome": "Influente", "estilo": "Pessoas e Comunicação", "desc": "Entusiasmado e otimista. Busca conexão social.", "cor": "orange", "tarefas": "Apresentações, Networking, Motivação."},
            "S": {"nome": "Estável", "estilo": "Colaboração e Persistência", "desc": "Paciente e leal. Busca harmonia e segurança.", "cor": "green", "tarefas": "Apoio operacional, Suporte, Processos."},
            "C": {"nome": "Conformidade", "estilo": "Precisão e Qualidade", "desc": "Analítico e detalhista. Busca lógica e regras.", "cor": "blue", "tarefas": "Auditoria, Análise de dados, Padronização."}
        }

        info = textos_disc.get(dominante, {"nome": "N/A", "estilo": "", "desc": "", "cor": "gray", "tarefas": ""})

        st.markdown(f"### Análise do Perfil: :{info['cor']}[{info['nome']}]")
        st.write(f"**Foco Principal:** {info['estilo']}")
        
        col_desc, col_tar = st.columns(2)
        with col_desc:
            st.info(info['desc'])
        with col_tar:
            st.warning(f"**Tarefas Sugeridas:**\n{info['tarefas']}")

        # 3. LEGENDA DETALHADA (Final da página)
        with st.expander("🔍 Legenda Geral DISC - Detalhada", expanded=False):
            textos_disc = {
                "D": {
                    "nome": "Dominante",
                    "estilo": "Resultados e Assertividade",
                    "desc": "Decidido e direto. Busca desafios, rapidez e liderança.",
                    "cargos": "Gerente, Líder de Projeto, Coordenador",
                    "tarefas_mais": "Tomada de decisão, Gestão de crises, Definir metas",
                    "tarefas_menos": "Atendimento de rotina, Processos detalhados, Documentação"
                },
                "I": {
                    "nome": "Influente",
                    "estilo": "Pessoas e Comunicação",
                    "desc": "Entusiasmado, sociável e persuasivo. Busca conexão e motivação do grupo.",
                    "cargos": "Marketing, Vendas, Comunicação, Treinamento",
                    "tarefas_mais": "Apresentações, Networking, Reuniões de equipe, Motivação",
                    "tarefas_menos": "Tarefas repetitivas, Processos rígidos, Detalhes técnicos"
                },
                "S": {
                    "nome": "Estável",
                    "estilo": "Colaboração e Persistência",
                    "desc": "Paciente, leal e confiável. Busca harmonia e segurança.",
                    "cargos": "Suporte, Administrativo, RH, Atendimento ao Cliente",
                    "tarefas_mais": "Suporte operacional, Atendimento, Organizar processos",
                    "tarefas_menos": "Mudanças constantes, Pressão por resultados rápidos, Competição intensa"
                },
                "C": {
                    "nome": "Conformidade",
                    "estilo": "Precisão e Qualidade",
                    "desc": "Analítico, detalhista e criterioso. Busca lógica, regras e perfeição.",
                    "cargos": "Auditoria, Contabilidade, TI, Qualidade",
                    "tarefas_mais": "Análise de dados, Relatórios, Controle de qualidade, Padronização",
                    "tarefas_menos": "Decisões rápidas sem dados, Interações sociais constantes, Ambiguidade"
                }
            }

            for key, info in textos_disc.items():
                st.markdown(f"### **{key} - {info['nome']}**")
                st.write(f"**Estilo de trabalho:** {info['estilo']}")
                st.write(f"**Descrição:** {info['desc']}")
                st.write(f"**Cargos mais compatíveis:** {info['cargos']}")
                st.write(f"**Atividades que combinam mais:** {info['tarefas_mais']}")
                st.write(f"**Atividades que combinam menos:** {info['tarefas_menos']}")
                st.markdown("---")

        
                
        # ============================================================
        # COMPATIBILIDADE CARGO × PERFIL DISC (APENAS MENSAGEM)
        # ============================================================

        st.markdown("### 🔹 Compatibilidade Cargo × Perfil DISC")

        cargo_atual = str(form.get("cargo", "")).lower()

        # Mapeamento de cargos por perfil dominante
        compatibilidade = {
            "D": ["gerente", "diretor", "coordenador", "lider", "gestor"],
            "I": ["vendas", "marketing", "comercial", "relacionamento", "comunicação"],
            "S": ["rh", "suporte", "atendimento", "administrativo", "operacional"],
            "C": ["contabilidade", "qualidade", "auditoria", "financeiro", "ti", "analista"]
        }

        cargos_compatíveis = compatibilidade.get(dominante, [])
        match = any(c in cargo_atual for c in cargos_compatíveis)

        # Exibição simplificada em métricas
        colA, colB = st.columns(2)
        colA.metric("Cargo Atual", form.get("cargo","N/A").title())
        colB.metric("Perfil Dominante", dominante if dominante else "N/A")

        # Mensagem direta sem gráfico
        if match:
            st.success(f"**Alta aderência:** O perfil **{dominante}** possui características naturais que favorecem o desempenho em cargos de **{cargo_atual.title()}**.")
        else:
            st.warning(f"**Ponto de Atenção:** O perfil **{dominante}** pode exigir um esforço maior de adaptação para as rotinas típicas de **{cargo_atual.title()}**.")

        # ============================================================
        # PERFIL DISC EXIGIDO PELAS ATIVIDADES
        # ============================================================

        st.markdown("### 🔹 Perfil DISC Exigido pelas Atividades")

        atividades_lista = [
            a.get("Atividade Descrita","")
            for a in form.get("atividades",[])
        ]

        atividades_texto = " ".join(atividades_lista).lower()

        compatibilidade_ativ = {

            "D": [
                "decisão","meta","resultado","liderar","negociar",
                "estratégia","direcionar","definir","priorizar"
            ],

            "I": [
                "apresentar","convencer","comunicar","clientes",
                "reunião","relacionamento","treinamento"
            ],

            "S": [
                "suporte","atender","organizar","rotina",
                "apoio","assistir","acompanhar","colaborar"
            ],

            "C": [
                "analisar","dados","relatório","planilha",
                "controle","auditar","conferir","classificar",
                "registrar","custos","informações","base",
                "indicadores","verificar","validar"
            ]

        }

        scores = {}

        for perfil, palavras in compatibilidade_ativ.items():

            pontos = sum(
                atividades_texto.count(p) for p in palavras
            )

            scores[perfil] = pontos

        perfil_exigido = max(scores, key=scores.get)

        # ============================================================
        # MÉTRICAS
        # ============================================================

        colA, colB, colC = st.columns(3)

        colA.metric("Perfil do Colaborador", dominante if dominante else "N/A")
        colB.metric("Perfil Exigido pelas Atividades", perfil_exigido)

        total_pontos = sum(scores.values())

        if total_pontos > 0:
            compat_percent = int((scores.get(dominante,0) / total_pontos) * 100)
        else:
            compat_percent = 0

        colC.metric("Compatibilidade", f"{compat_percent}%")

        # ============================================================
        # MENSAGEM PRINCIPAL
        # ============================================================

        if perfil_exigido == dominante:

            st.success(
                f"Alta aderência: As atividades indicam um perfil **{perfil_exigido}**, compatível com o perfil do colaborador."
            )

        else:

            st.warning(
                f"As atividades indicam um perfil **{perfil_exigido}**, enquanto o colaborador apresenta perfil **{dominante}**."
            )

        # ============================================================
        # ATIVIDADES QUE EXIGEM ADAPTAÇÃO
        # ============================================================

        atividades_compativeis = compatibilidade_ativ.get(perfil_exigido, [])

        atividades_desvio = []

        for ativ in atividades_lista:

            texto = str(ativ).lower()

            if not any(p in texto for p in atividades_compativeis):
                atividades_desvio.append(ativ)


        ranking_atividades = []

        for ativ in atividades_lista:

            texto = str(ativ).lower()

            if not texto.strip():
                continue

            score = sum(p in texto for p in compatibilidade_ativ.get(dominante, []))

            ranking_atividades.append((score, ativ))


        ranking_atividades.sort(key=lambda x: x[0])


        if ranking_atividades:

            st.markdown("#### ⚠ Lista das principais dificuldades de adaptação")

            limite = min(3, len(ranking_atividades))

            for score, atividade in ranking_atividades[:limite]:
                st.write("•", atividade)


import streamlit as st
import pandas as pd
import os
import json
import sys

# ============================================================
# CONFIGURAÇÃO DE DIRETÓRIO E CARREGAMENTO
# ============================================================

# Define o diretório base e a pasta de dados
base_dir = os.path.dirname(os.path.abspath(__file__))
dados_dir = os.path.join(base_dir, "dados")
os.makedirs(dados_dir, exist_ok=True)

# Função para carregar todos os JSONs da pasta 'dados'
def carregar_todos_formularios():
    lista_formularios = []
    if os.path.exists(dados_dir):
        for nome_arquivo in os.listdir(dados_dir):
            if nome_arquivo.endswith(".json"):
                caminho_completo = os.path.join(dados_dir, nome_arquivo)
                try:
                    with open(caminho_completo, "r", encoding="utf-8") as f:
                        dados = json.load(f)
                        if isinstance(dados, dict):
                            lista_formularios.append(dados)
                except Exception:
                    continue
    return lista_formularios

# Inicializa o estado da sessão com os dados carregados
if "formularios" not in st.session_state:
    st.session_state["formularios"] = carregar_todos_formularios()

# --- BLOCO DE CSS PARA OCULTAÇÃO ---
if st.query_params.get("page") == "formulario":
    st.markdown("""
    <style>
        [data-testid="stSidebar"] {display: none !important;}
        #MainMenu, footer, header {visibility: hidden !important;}
    </style>
    """, unsafe_allow_html=True)


# =========================================================
# 1. FUNÇÕES DE SUPORTE
# =========================================================
def preparar_df(chave_json, colunas, fonte_local, linhas_padrao=15):
    if not isinstance(fonte_local, dict): fonte_local = {}
    dados = fonte_local.get(chave_json, [])
    if dados and isinstance(dados, list):
        df = pd.DataFrame(dados)
        for col in colunas:
            if col not in df.columns: df[col] = ""
        return df[colunas]
    return pd.DataFrame({col: [""] * linhas_padrao for col in colunas})

lista_frequencia = ["", "DVD", "D", "S", "Q", "M", "T", "A"]
lista_horas = [f"{i} h" for i in range(25)]
lista_minutos = [f"{i} min" for i in range(0, 60, 5)]

col_atv = ["Atividade", "Frequência", "Horas", "Minutos"]
# Altere para:
col_dif = ["Dificuldade/Bloqueio", "Setor/Parceiro Envolvido", "Frequência", "Horas", "Minutos"]
col_sug = ["Sugestão de Melhoria", "Impacto Esperado", "Frequência", "Horas", "Minutos"]


# =========================================================
# 🎭 CAPA DE INVISIBILIDADE (NÃO QUEBRA O RASCUNHO)
# =========================================================

# Criamos um lugar no app que pode ser "esvaziado"
area_do_formulario = st.container()

# Se NÃO estivermos na página do formulário, a gente limpa a área visual
# mas deixa o código das 3000 linhas rodar "em silêncio" para o rascunho
if st.session_state.get("pagina") != "formulario":
    area_do_formulario.empty() 

# Agora, para o Título e as mensagens iniciais, usamos o 'with'
with area_do_formulario:
    resgate = st.session_state.get("rascunho_atual", {})
    nome_titulo = resgate.get("colaborador", "Novo Formulário")

    st.title("📋 Formulário de Acompanhamento")

    if nome_titulo != "Novo Formulário":
        st.info(f"✨ **Editando Rascunho de:** {nome_titulo}")
    else:
        st.success("📝 **Criando Novo Registro**")

    st.markdown("---")

# ABAIXO SEGUEM AS 3000 LINHAS SEM INDENTAÇÃO
# O rascunho continuará funcionando porque o código está sendo lido,
# mas o 'area_do_formulario.empty()' lá em cima ajuda a limpar o topo.

# =========================================================
# Perguntas DISC
# =========================================================
perguntas_disc = [
    "Quando surge um problema inesperado: (A) Age rápido | (B) Comunica a todos | (C) Analisa riscos | (D) Segue processo",
    "Em situações de pressão: (A) Foca no resultado | (B) Mantém o otimismo | (C) Mantém a calma | (D) Busca precisão",
    "Ao receber tarefa difícil: (A) Aceita o desafio | (B) Busca ajuda social | (C) Planeja passos | (D) Estuda as regras",
    "No trabalho em equipe: (A) Lidera o grupo | (B) Motiva os colegas | (C) Apoia os outros | (D) Organiza as tarefas",
    "Em reuniões: (A) Vai direto ao ponto | (B) Interage e brinca | (C) Escuta mais | (D) Anota detalhes",
    "Ao lidar com conflitos: (A) Enfrenta direto | (B) Tenta apaziguar | (C) Evita o confronto | (D) Usa lógica e fatos",
    "Seu ritmo de trabalho: (A) Rápido/Impaciente | (B) Rápido/Entusiasmado | (C) Calmo/Constantemente | (D) Metódico/Cauteloso",
    "Prefere tarefas: (A) Desafiadoras | (B) Variadas e sociais | (C) Rotineiras e seguras | (D) Técnicas e detalhadas",
    "Seu foco principal: (A) Resultados | (B) Relacionamentos | (C) Estabilidade | (D) Qualidade e Processos",
    "Ao decidir, você é: (A) Decidido e firme | (B) Impulsivo e intuitivo | (C) Cuidadoso e lento | (D) Lógico e analítico",
    "Confia mais em: (A) Sua intuição | (B) Opinião alheia | (C) Experiência passada | (D) Dados e provas",
    "Prefere decisões: (A) Independentes | (B) Em grupo | (C) Consensuais | (D) Baseadas em normas",
    "Estilo de organização: (A) Prático | (B) Criativo/Bagunçado | (C) Tradicional | (D) Muito organizado",
    "Lida melhor com: (A) Mudanças rápidas | (B) Novas ideias | (C) Rotinas claras | (D) Regras rígidas",
    "Prefere trabalhar: (A) Sozinho/Comando | (B) Ambiente festivo | (C) Ambiente tranquilo | (D) Ambiente silencioso",
    "Seu ponto forte: (A) Coragem | (B) Comunicação | (C) Paciência | (D) Organização",
    "Você se considera: (A) Dominante | (B) Influente | (C) Estável | (D) Conforme/Analítico",
    "Se motiva por: (A) Poder/Bônus | (B) Reconhecimento | (C) Segurança/Paz | (D) Conhecimento Técnico",
    "Reação a cobranças: (A) Mais esforço | (B) Desculpas criativas | (C) Ansiedade | (D) Argumentos técnicos",
    "Ambiente ideal: (A) Competitivo | (B) Amigável | (C) Previsível | (D) Disciplinado",
    "Ao lidar com feedback: (A) Aceita e ajusta | (B) Comenta e debate | (C) Analisa e planeja | (D) Segue regras",
    "Como prefere aprender: (A) Fazendo | (B) Interagindo | (C) Observando | (D) Estudando materiais",
    "Gestão de tempo: (A) Prioriza resultados | (B) Mantém relações | (C) Planeja com cuidado | (D) Segue processos",
    "Como se comunica: (A) Direto e objetivo | (B) Amigável e motivador | (C) Calmo e ponderado | (D) Técnico e detalhista"
]


# =========================================================
# 👤 DADOS DE IDENTIFICAÇÃO (AJUSTADO PARA 5 TABELAS)
# =========================================================
st.subheader("👤 Dados de Identificação")

st.write("DEBUG - O que tem no state agora:", st.session_state.get("f_cargo"))

fonte = st.session_state.get("dados_oficiais", {})
col1, col2 = st.columns(2)

with col1:
    # Mostra os nomes encontrados no GitHub
    rascunhos_dict = st.session_state.get("rascunhos", {})
    nomes_disponiveis = list(rascunhos_dict.keys())
    st.write(f"🗂️ Rascunhos na Nuvem: **{', '.join(nomes_disponiveis) if nomes_disponiveis else 'Nenhum'}**")

    v = st.session_state.get("v_tab", 0)
    nome_f = st.text_input(
        "Nome do colaborador",
        value=st.session_state.get("f_nome_v2") or fonte.get("nome", ""),
        key=f"f_nome_{v}"
    )

    if st.button("📥 Carregar Rascunho", key="btn_carregar_rascunho_v3"):
        if nome_f:
            atualizar_rascunhos_do_github() 
            rascunhos_dict = st.session_state.get("rascunhos", {})
            
            nome_busca = nome_f.strip().upper()
            rascunho = rascunhos_dict.get(nome_busca)
            
            if rascunho:
                # 1. SALVA O ESTADO GLOBAL
                st.session_state["rascunho"] = rascunho  # <--- CRUCIAL para o motor de tabelas
                st.session_state["f_nome_v2"] = nome_busca
                st.session_state["v_tab"] = st.session_state.get("v_tab", 0) + 1
                
                # 2. DADOS BÁSICOS (Ajustado para a estrutura do seu JSON)
                cp = rascunho.get("campos", {})
                
                # Mapeamento: "Chave_do_Widget": "Chave_dentro_do_JSON_campos"
                campos_map = {
                    "f_cargo": "cargo",
                    "f_depto": "dep",         # No JSON está "dep", não "departamento"
                    "f_setor": "setor",
                    "f_chefe": "chefe",
                    "f_unidade": "unidade",
                    "f_esc": "escolaridade",
                    "f_dev": "devolver_em",    # No JSON está "devolver_em"
                    "f_cursos_area": "cursos",
                    "f_obj_area": "objetivo"
                }
                
                for key_ui, key_json in campos_map.items():
                    valor = cp.get(key_json, "")
                    st.session_state[key_ui] = valor
                    st.session_state[f"{key_ui}_v2"] = valor

                # 3. DISC - SINCRONIZAÇÃO TOTAL
                disc_salvo = rascunho.get("disc", {})
                if disc_salvo:
                    for i in range(24):
                        chave_json = str(i)
                        # Sincroniza a chave do rádio que você usa no loop do DISC
                        # Importante: a chave aqui deve ser idêntica à definida no st.radio
                        v = st.session_state["v_tab"]
                        st.session_state[f"disc_radio_{i}_{v}"] = disc_salvo.get(chave_json)


                # --- 4. TABELAS - SINCRONIZAÇÃO ---
                tabelas_salvas = rascunho.get("tabelas", {})
                if tabelas_salvas:
                    # Lista das chaves de tabela que você tem no JSON
                    chaves_tabelas = ["alta", "normal", "baixa", "dificuldades", "sugestoes"]
                    
                    for t_key in chaves_tabelas:
                        dados = tabelas_salvas.get(t_key, [])
                        if dados:
                            # Converte a lista de dicionários do JSON em um DataFrame
                            df_carregado = pd.DataFrame(dados)
                            
                            # SALVA NO SESSION_STATE 
                            # Importante: A chave deve ser EXATAMENTE a 'key' que você usa no st.data_editor
                            # Se o seu data_editor usa key=f"editor_{t_key}_{v}", faça igual:
                            v = st.session_state["v_tab"]
                            st.session_state[f"editor_{t_key}_{v}"] = df_carregado

                st.success(f"✅ Rascunho e DISC de {nome_busca} carregados!")
                st.rerun()
            else:
                st.error(f"⚠️ Rascunho de '{nome_busca}' não encontrado.")
        else:
            st.warning("⚠️ Digite um nome antes de carregar.")

with col2:
    cargo_f = st.text_input("Cargo", value=st.session_state.get("f_cargo_v2") or fonte.get("cargo", ""), key="f_cargo")
    depto_f = st.text_input("Departamento", value=st.session_state.get("f_depto_v2") or fonte.get("departamento", ""), key="f_depto")
    esc_f = st.text_input("Escolaridade", value=st.session_state.get("f_esc_v2") or fonte.get("escolaridade", ""), key="f_esc")
    setor_f = st.text_input("Setor", value=st.session_state.get("f_setor_v2") or fonte.get("setor", ""), key="f_setor")
    chefe_f = st.text_input("Chefe imediato", value=st.session_state.get("f_chefe_v2") or fonte.get("chefe", ""), key="f_chefe")
    unidade_f = st.text_input("Empresa / Unidade", value=st.session_state.get("f_unidade_v2") or fonte.get("unidade", ""), key="f_unidade")
    dev_f = st.text_input("Devolver preenchido em", value=st.session_state.get("f_dev_v2") or fonte.get("devolucao", ""), key="f_dev")

cursos_f = st.text_area("Cursos Obrigatórios e Diferenciais", value=st.session_state.get("f_cursos_v2") or fonte.get("cursos", ""), key="f_cursos_area")
obj_f = st.text_area("Em que consiste seu Trabalho e qual seu Principal Objetivo", value=st.session_state.get("f_obj_v2") or fonte.get("objetivo", ""), key="f_obj_area")

    

# =========================================================
# 5. TABELAS DE TAREFAS (COM FUNÇÃO DE SUPORTE INTEGRADA)
# =========================================================
st.markdown("---")

st.subheader("📋 Tabelas") # Título médio

# --- INÍCIO DAS LEGENDAS ---
col_leg1, col_leg2 = st.columns(2)

with col_leg1:
    st.info("""
    **📋 LEGENDA DE FREQUÊNCIA:**
    * **DVD**: Diário Várias Vezes
    * **D**: Diário | **S**: Semanal
    * **Q**: Quinzenal | **M**: Mensal
    * **T**: Trimestral | **A**: Anual
    """)

with col_leg2:
    st.warning("""
    **⏱️ COMO REGISTRAR O TEMPO:**
    * **Horas e Minutos**: Selecione o valor em cada coluna.
    * **Menos de 1 hora?**: Selecione **0 h** e o tempo real em minutos.
    * **Não se aplica?**: Selecione **0 h** e **0 min** em ambos.
    """)
# --- FIM DAS LEGENDAS ---

# --- FUNÇÃO AUXILIAR (Garante que o editor tenha linhas suficientes) ---
def garantir_15_linhas(df, colunas):
    if df is None or df.empty:
        df = pd.DataFrame(columns=colunas)
    for col in colunas:
        if col not in df.columns: df[col] = ""
    while len(df) < 15:
        df.loc[len(df)] = [""] * len(colunas)
    return df.head(15)

# 1. Configurações e rascunho
lista_frequencia = ["", "DVD", "D", "S", "Q", "M", "T", "A"]
lista_horas = [f"{i} h" for i in range(25)]
lista_minutos = [f"{i} min" for i in range(0, 60, 5)]

if "rascunho_atual" not in st.session_state:
    st.session_state["rascunho_atual"] = {}

rascunho = st.session_state.get("rascunho", {})
v_layout = st.session_state.get("v_tab", 0)

# 2. Definição da função de renderização
def gerar_tabela_final(titulo, chave_json, col_principal, col_extra=None, label_extra=None):
    st.subheader(titulo)
    dict_tabelas = rascunho.get("tabelas", {}) if isinstance(rascunho, dict) else {}
    dados_salvos = dict_tabelas.get(chave_json, [])
    
    colunas = [col_principal, "Horas", "Minutos", "Frequência"]
    if col_extra: 
        colunas.insert(1, col_extra)
    
    # Chama a função que estava dando NameError
    df_base = garantir_15_linhas(pd.DataFrame(dados_salvos), colunas)
    
    config_tab = {
        col_principal: st.column_config.TextColumn("Descrição", width="large"),
        "Frequência": st.column_config.SelectboxColumn("Frequência", options=lista_frequencia, width="small"),
        "Horas": st.column_config.SelectboxColumn("Horas", options=lista_horas, width="small"),
        "Minutos": st.column_config.SelectboxColumn("Minutos", options=lista_minutos, width="small"),
    }
    if col_extra: 
        config_tab[col_extra] = st.column_config.TextColumn(label_extra, width="medium")

    return st.data_editor(
        df_base, 
        key=f"editor_{chave_json}_v{v_layout}", 
        column_config=config_tab, 
        use_container_width=True, 
        num_rows="fixed"
    )

# 3. Chamadas das Tabelas (Sincronização Total com o JSON)
# Nota: O terceiro parâmetro deve ser EXATAMENTE a chave do JSON (ex: "Atividade")

e_alta = gerar_tabela_final("🚀 Atividades de Alta Complexidade", "alta", "Atividade")

e_normal = gerar_tabela_final("📋 Atividades de Complexidade Normal", "normal", "Atividade")

e_baixa = gerar_tabela_final("⏳ Atividades de Baixa Complexidade", "baixa", "Atividade")

# Aqui mudamos de "Setor/Parceiro Envolvido" para "Setor Envolvido"
e_dif = gerar_tabela_final("⚠️ Dificuldades e Bloqueios", "dificuldades", "Dificuldade", "Setor Envolvido", "Setor Envolvido")

# Aqui mudamos de "Impacto" para "Impacto Esperado"
e_sug = gerar_tabela_final("💡 Sugestões de Melhoria", "sugestoes", "Sugestão", "Impacto Esperado", "Impacto Esperado")


# =========================================================
# 📊 7. QUESTIONÁRIO DISC (SINCRONIZADO COM O JSON)
# =========================================================
st.markdown("---")
st.subheader("📊 Questionário")

v = st.session_state.get("v_tab", 0) 

# 1. BUSCA O RASCUNHO NO LUGAR CERTO (Onde o JSON que você mostrou reside)
# O seu JSON mostra que o DISC está dentro de "rascunho" -> "disc"
rascunho_disc = st.session_state.get("rascunho", {}).get("disc", {})

respostas_disc_atual = {}

for i, pergunta in enumerate(perguntas_disc):
    # 2. BUSCA PELA CHAVE DO JSON (O seu JSON usa apenas o número como string)
    chave_json = str(i)
    letra_salva = rascunho_disc.get(chave_json)
    
    # 3. DEFINE O ÍNDICE (A=0, B=1, C=2, D=3)
    opcoes = ["A", "B", "C", "D"]
    # Se letra_salva for None ou null no JSON, o index será None (fica desmarcado)
    idx_selecionado = opcoes.index(letra_salva) if letra_salva in opcoes else None
    
    # 4. O WIDGET
    escolha = st.radio(
        f"**{i+1}.** {pergunta}",
        options=opcoes,
        index=idx_selecionado,
        key=f"disc_radio_{i}_{v}",
        horizontal=True
    )
    
    # 5. GUARDA PARA SALVAR DEPOIS (Mantendo o padrão de string do JSON)
    respostas_disc_atual[chave_json] = escolha

# 6. ATUALIZA O RASCUNHO GLOBAL NA HORA
# Verifica se o rascunho existe; se não, inicializa para evitar o KeyError
if "rascunho" not in st.session_state:
    st.session_state["rascunho"] = {}

# Garante que a sub-chave "disc" também exista
if "disc" not in st.session_state["rascunho"]:
    st.session_state["rascunho"]["disc"] = {}

# Agora sim, salva as respostas sem risco de quebrar
st.session_state["rascunho"]["disc"] = respostas_disc_atual



# =========================================================
# 6. VALIDAÇÃO UNIFICADA (TABELAS, DISC E CABEÇALHO)
# =========================================================
st.markdown("---")
st.subheader("✅ Status de Validação do Formulário")

pendencias = []

# --- 1. VALIDAÇÃO DE CABEÇALHO ---
campos_id = {
    "Nome": nome_f, "Cargo": cargo_f, "Departamento": depto_f,
    "Escolaridade": esc_f, "Setor": setor_f, "Chefe Imediato": chefe_f,
    "Empresa/Unidade": unidade_f, "Devolver em": dev_f,
    "Cursos": cursos_f, "Objetivo": obj_f
}
for campo, valor in campos_id.items():
    if not valor or str(valor).strip() == "":
        pendencias.append(f"Identificação: O campo **{campo}** está vazio.")

# --- 2. VALIDAÇÃO DAS TABELAS (RIGOR TOTAL) ---
dict_tabelas = {
    "Alta Complexidade": e_alta, 
    "Complexidade Normal": e_normal,
    "Baixa Complexidade": e_baixa, 
    "Dificuldades": e_dif,
    "Sugestões e Melhorias": e_sug
}

regras_colunas = {
    "Alta Complexidade": "Atividade", 
    "Complexidade Normal": "Atividade",
    "Baixa Complexidade": "Atividade", 
    "Dificuldades": "Dificuldade",
    "Sugestões e Melhorias": "Sugestão"
}

for nome_tab, df_validar in dict_tabelas.items():
    col_alvo = regras_colunas.get(nome_tab)
    
    if df_validar is not None and col_alvo in df_validar.columns:
        # Identifica linhas onde a descrição foi preenchida
        linhas_ativas = df_validar[df_validar[col_alvo].astype(str).str.strip() != ""]
        
        if len(linhas_ativas) == 0:
            pendencias.append(f"⚠️ A tabela **{nome_tab}** está totalmente vazia. Preencha pelo menos 1 item.")
        else:
            for i, row in linhas_ativas.iterrows():
                # Extração limpa dos valores
                h_str = str(row.get("Horas", "")).strip()
                m_str = str(row.get("Minutos", "")).strip()
                freq = str(row.get("Frequência", "")).strip()
                
                # Validação de Horas
                if h_str == "":
                    pendencias.append(f"❌ {nome_tab} (Linha {i+1}): Falta selecionar as **Horas**.")
                
                # Validação de Minutos
                if m_str == "":
                    pendencias.append(f"❌ {nome_tab} (Linha {i+1}): Falta selecionar os **Minutos**.")
                
                # Validação de Frequência
                if freq == "":
                    pendencias.append(f"❌ {nome_tab} (Linha {i+1}): Falta selecionar a **Frequência**.")
                
                # Validação extra para colunas específicas (Impacto / Setor)
                if nome_tab == "Dificuldades":
                    if str(row.get("Setor Envolvido", "")).strip() == "":
                        pendencias.append(f"❌ {nome_tab} (Linha {i+1}): Informe o **Setor Envolvido**.")
                
                if nome_tab == "Sugestões e Melhorias":
                    if str(row.get("Impacto Esperado", "")).strip() == "":
                        pendencias.append(f"❌ {nome_tab} (Linha {i+1}): Informe o **Impacto Esperado**.")

# --- 3. VALIDAÇÃO DO DISC ---
respostas_vazias = [k for k, v in respostas_disc_atual.items() if v is None]
if len(respostas_vazias) > 0:
    pendencias.append(f"Questionário: Faltam responder **{len(respostas_vazias)} questões**.")

# --- EXIBIÇÃO FINAL DO STATUS ---
if pendencias:
    st.warning(f"⚠️ **Existem {len(pendencias)} pendências obrigatórias:**")
    for p in pendencias:
        st.write(f"• {p}")
    st.session_state["confirmacao_final"] = False
else:
    st.success("🎉 **Perfeito! Tudo preenchido corretamente. O envio está liberado.**")

# =========================================================
# 🚀 4. BOTÃO DE ENVIO E SALVAMENTO REAL (VERSÃO FINAL)
# =========================================================

# Centralizando o botão para dar mais destaque
col_btn, _ = st.columns([2, 1])

with col_btn:
    if st.button("🚀 FINALIZAR E ENVIAR FORMULÁRIO", type="primary", use_container_width=True):
        if pendencias:
            st.error("❌ Corrija as pendências listadas acima antes de enviar.")
            st.stop()
        
        # Sistema de confirmação dupla
        if not st.session_state.get("confirmacao_final", False):
            st.warning(f"⚠️ **{nome_f}**, clique novamente para confirmar o envio definitivo.")
            st.session_state["confirmacao_final"] = True
            st.stop()

        try:
            from datetime import datetime
            import json

            def preparar_dados(df):
                if df is None or df.empty: return []
                # Pega dinamicamente a coluna de Descrição (Atividade, Dificuldade ou Sugestão)
                col_principal = df.columns[0] 
                return df[df[col_principal].astype(str).str.strip() != ""].to_dict("records")

            # 1. EXTRAÇÃO FORÇADA (Garante que os dados saiam do rádio e entrem no código)
            dados_disc_final = {}
            for i in range(24):
                # Tentamos todas as variações possíveis de nomes que você pode ter usado
                valor = st.session_state.get(f"q_{i}") or st.session_state.get(f"q{i}") or st.session_state.get(f"p{i}") or ""
                dados_disc_final[str(i)] = valor

            payload = {
                "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                "colaborador": nome_f,
                "campos": {
                    "cargo": cargo_f,
                    "departamento": depto_f,
                    "setor": setor_f,
                    "chefe": chefe_f,
                    "unidade": unidade_f,
                    "escolaridade": esc_f,
                    "devolver_em": dev_f,
                    "cursos": cursos_f,
                    "objetivo": obj_f
                },
                "tabelas": {
                    "alta": preparar_dados(e_alta),
                    "normal": preparar_dados(e_normal),
                    "baixa": preparar_dados(e_baixa),
                    "dificuldades": preparar_dados(e_dif),
                    "sugestoes": preparar_dados(e_sug)
                },
                "disc": st.session_state.get("rascunho", {}).get("disc", {})
            }




            nome_arquivo = f"{nome_f.replace(' ', '_').upper()}.json"
            
            with st.spinner("Sincronizando..."):
                sucesso = salvar_no_github(payload, nome_arquivo, pasta="dados")

                if sucesso:
                    
                    st.success(f"✅ Formulário de {nome_f} enviado com sucesso!")
                    enviado = enviar_para_sheets(payload)

                    if enviado:
                        st.toast("📊 Enviado para Google Sheets!")
                    else:
                        st.warning("⚠️ Salvou, mas não enviou para o Sheets")



                    # Limpa os estados de controle
                    st.session_state["confirmacao_final"] = False
                    # Opcional: st.session_state["rascunho_atual"] = {} 
                else:
                    st.error("⚠️ O GitHub não respondeu. Mas não se preocupe! Baixe o arquivo abaixo e envie por e-mail/WhatsApp.")

                # Botão de download sempre visível após tentativa de envio (Backup)
                st.download_button(
                    label="📥 Baixar Cópia de Segurança (JSON)",
                    data=json.dumps(payload, indent=4, ensure_ascii=False),
                    file_name=nome_arquivo,
                    mime="application/json",
                    use_container_width=True
                )

        except Exception as e:
            st.error(f"❌ Erro crítico ao processar o envio: {e}")

# --- VISUALIZAÇÃO ---
if st.session_state.get("pagina") == "visualizar":
    st.title("👁️ Visualização de Registros")
    
    # 1. Carrega os dados frescos do disco
    lista_de_arquivos = carregar_todos_formularios()
    
    # 2. Se a sua função carregar_todos_formularios() já retorna a lista, 
    # apenas certifique-se de que não estamos adicionando isso ao session_state de forma acumulativa.
    if not lista_de_arquivos:
        st.warning("⚠️ Nenhum formulário encontrado.")
    else:
        # Mostra o total para conferência
        st.success(f"Foram encontrados {len(lista_de_arquivos)} formulários.")
        
        # 3. Exibição limpa
        for idx, form in enumerate(lista_de_arquivos, 1):
            nome_exibir = str(form.get('nome', f'Colaborador {idx}')).upper()
            
           
            with st.expander(f"👤 FORMULÁRIO DE: {nome_exibir}", expanded=True):               
                
            
                            
                # 1. Cabeçalho Completo
                st.subheader("📝 Informações de Identificação")
                col1, col2 = st.columns(2)
                col1.write(f"**Data de Envio:** {form.get('data_envio', 'N/A')}")
                col2.write(f"**Devolver em:** {form.get('devolucao', 'N/A')}")
                
                col_a, col_b = st.columns(2)
                col_a.write(f"**Setor:** {form.get('setor', 'N/A')}")
                col_b.write(f"**Departamento:** {form.get('departamento', 'N/A')}")
                col_a.write(f"**Cargo:** {form.get('cargo', 'N/A')}")
                col_b.write(f"**Chefe Imediato:** {form.get('chefe', 'N/A')}")
                col_a.write(f"**Empresa/Unidade:** {form.get('empresa', 'N/A')}")
                col_b.write(f"**Escolaridade:** {form.get('escolaridade', 'N/A')}")
                
                st.subheader("🎓 Cursos Obrigatórios ou Diferenciais")

                st.info(
                    form.get("cursos", "Não informado")
                )

                st.subheader("🎯 Trabalho e Principal Objetivo")

                st.info(
                    form.get("objetivo", "Não informado")
                )
                
                # 2. Tabelas Dinâmicas
                secoes = {
                    "atividades": "📋 Atividades Executadas",
                    "dificuldades": "⚠️ Dificuldades e Bloqueios",
                    "sugestoes": "💡 Sugestões de Melhoria"
                }
                
                for chave, titulo in secoes.items():
                    st.markdown("---")
                    st.subheader(titulo)
                    if chave in form and form[chave]:
                        df = pd.DataFrame(form[chave])
                        df = df.replace("", None).dropna(how='all')
                        if not df.empty:
                            st.table(df)
                        else:
                            st.write("Nenhum dado preenchido nesta seção.")
                    else:
                        st.write("Seção não encontrada ou vazia.")
                
                # 3. Questionário DISC (Exibição Completa e Legível)
                st.markdown("---")
                st.subheader("📊 Avaliação DISC (Perguntas e Respostas)")
                
                for i, pergunta in enumerate(perguntas_disc, 1):
                    valor_resposta = form.get("disc", {}).get(f"disc_{i}", "Não respondido")
                    st.write(f"**{i}. {pergunta}**")
                    st.info(f"Resposta selecionada: **{valor_resposta}**")
                    st.markdown("---")

                # -------------------------------------------------
                # BOTÕES DE EXPORTAÇÃO
                # -------------------------------------------------
                if st.session_state.get("usuario_logado") == "Luciano 123":
                    st.markdown("---")
                    st.subheader("⚙️ Painel de Exportação")

                    col1, col2 = st.columns(2)

                    # Preparação do nome do arquivo
                    nome = form.get("nome", "Colaborador")
                    data = form.get("data_envio", "")
                    
                    # Limpeza para evitar caracteres proibidos em nomes de arquivos
                    nome_clean = nome.replace(" ", "_")
                    data_clean = str(data).replace("/", "").replace(":", "").replace(" ", "_")
                    nome_arquivo = f"Relatorio_{nome_clean}_{data_clean}"

                    with col1:
                        st.download_button(
                            label="📄 Baixar Word",
                            data=gerar_word(form),
                            file_name=f"{nome_arquivo}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"btn_word_{idx}" # Key estável usando o índice do loop
                        )

                    with col2:
                        st.download_button(
                            label="📑 Baixar PDF",
                            data=gerar_pdf(form),
                            file_name=f"{nome_arquivo}.pdf",
                            mime="application/pdf",
                            key=f"btn_pdf_{idx}" # Key estável usando o índice do loop
                        )



        st.markdown("---")
        st.subheader("🗑️ Excluir formulário específico")

        # Lista os arquivos
        arquivos_json = [f for f in os.listdir(dados_dir) if f.endswith(".json")]

        if arquivos_json:

            # Criar lista com nome do colaborador
            opcoes = []

            for arquivo in arquivos_json:
                caminho = os.path.join(dados_dir, arquivo)

                with open(caminho, "r", encoding="utf-8") as f:
                    try:
                        dados = json.load(f)

                        if isinstance(dados, dict):
                            nome = dados.get("nome", "Colaborador")
                        else:
                            nome = "Registro inválido"

                    except:
                        nome = "Arquivo corrompido"

                opcoes.append((arquivo, nome))

            # Mostrar opções
            nomes_para_select = [f"{nome} ({arquivo})" for arquivo, nome in opcoes]

            escolha = st.selectbox(
                "Selecione o formulário que deseja excluir:",
                nomes_para_select
            )

            if st.button("❌ Excluir formulário selecionado"):

                arquivo_escolhido = opcoes[nomes_para_select.index(escolha)][0]

                os.remove(os.path.join(dados_dir, arquivo_escolhido))

                st.success("✅ Formulário excluído com sucesso!")
                st.rerun()

        else:
            st.info("Nenhum formulário salvo.")

# ============================================================
# CALCULAR CARGA HORÁRIA
# ============================================================

def calcular_carga(atividades):
    total_min = 0
    for at in atividades:
        try:
            tempo = float(at.get("tempo","0"))
        except:
            tempo = 0
        freq = at.get("frequencia","semanal").lower()
        if freq == "diaria":
            total_min += tempo * 5
        elif freq == "mensal":
            total_min += tempo / 4
        else:
            total_min += tempo
    horas = total_min / 60
    status = "Adequado"
    if horas > 44: status = "Sobrecarga"
    elif horas < 30: status = "Subutilização"
    return round(horas,2), status

# ============================================================
# GERAR ATIVIDADES IDEAIS (GPT)
# ============================================================

def gerar_atividades_ideais(cargo, setor, client=None):
    if client is None:
        return [{
            "nome_atividade": "Atividade de exemplo",
            "descricao": "Descrição de exemplo",
            "frequencia_ideal": "semanal",
            "tempo_medio_minutos": 60,
            "justificativa_tecnica": "Exemplo"
        }]
    
    prompt = f"""
    Gere 12 atividades ideais para:
    Cargo: {cargo}
    Setor: {setor}
    Para cada atividade informe:
      - nome_atividade
      - descricao
      - frequencia_ideal
      - tempo_medio_minutos
      - justificativa_tecnica
    Responda SOMENTE JSON válido.
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.3
        )
        
        # 1. Primeiro carregamos o conteúdo em uma variável
        dados_carregados = json.loads(response.choices[0].message.content)

        # 2. Agora injetamos o DISC na memória antes de sair da função
        if isinstance(dados_carregados, dict) and "disc" in dados_carregados:
            st.session_state["respostas_disc_fix"] = {
                str(k): v for k, v in dados_carregados["disc"].items()
            }
        
        # 3. SÓ AGORA damos o return com os dados prontos
        return dados_carregados

    except Exception as e:
        # Se der erro, retorna o padrão
        return [{
            "nome_atividade": "Atividade de exemplo",
            "descricao": "Descrição de exemplo",
            "frequencia_ideal": "semanal",
            "tempo_medio_minutos": 60,
            "justificativa_tecnica": "Exemplo"
        }]

# ============================================================
# COMPARAÇÃO SEMÂNTICA
# ============================================================

def comparar_semanticamente(reais, ideais, client=None):
    if client is None:
        return {"score_aderencia":0,"tempo_gap_medio_percentual":0,"atividades_desvio":[]}

    prompt = f"""
    Compare semanticamente:
    Atividades reais: {reais}
    Atividades ideais: {ideais}
    Retorne JSON com:
      - score_aderencia (0-100)
      - tempo_gap_medio_percentual
      - atividades_desvio
    """
    try:
        r = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.2
        )
        return json.loads(r.choices[0].message.content)
    except:
        return {"score_aderencia":0,"tempo_gap_medio_percentual":0,"atividades_desvio":[]}

# ============================================================
# CLASSIFICAR DIFICULDADES
# ============================================================

def classificar_dificuldades_gpt(dificuldades, client=None):
    if client is None:
        return {}
    
    prompt = f"""
    Classifique semanticamente as dificuldades abaixo em:
    - Processo
    - Tempo
    - Comunicação
    - Estrutura
    - Liderança
    - Sistema
    Retorne JSON com contagem por categoria.
    Dificuldades: {dificuldades}
    """
    try:
        r = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}],
            temperature=0.2
        )
        return json.loads(r.choices[0].message.content)
    except:
        return {}

# ============================================================
# ÍNDICE GERAL DO CARGO
# ============================================================

def indice_geral(score_aderencia, score_disc, status_carga):
    fator_carga = 100
    if status_carga == "Sobrecarga": fator_carga = 70
    elif status_carga == "Subutilização": fator_carga = 75
    return round(mean([score_aderencia, score_disc, fator_carga]),2)

# ============================================================
# MOTOR PRINCIPAL COMPLETO – ANÁLISE CORPORATIVA
# ============================================================

def gerar_analise_corporativa(dados, client=None):
    """
    Gera análise completa de um colaborador com base em:
    - Atividades reais
    - Perfil DISC
    - Dificuldades
    Retorna:
    - parecer (texto)
    - indicadores (dict)
    """
    # 1️⃣ Atividades ideais
    ideais = gerar_atividades_ideais(dados["cargo"], dados["setor"], client)

    # 2️⃣ Comparação semântica (reais x ideais)
    comparacao = comparar_semanticamente(dados["atividades"], ideais, client)

    # 3️⃣ Carga horária
    horas, status_carga = calcular_carga(dados["atividades"])

    # 4️⃣ Score DISC
    disc_score = score_disc(dados["disc"])

    # 5️⃣ Classificação de dificuldades
    dificuldades_classificadas = classificar_dificuldades_gpt(dados["dificuldades"], client)

    # 6️⃣ Score de aderência
    score_aderencia = comparacao.get("score_aderencia",0)

    # 7️⃣ Índice geral
    indice = indice_geral(score_aderencia, disc_score, status_carga)

    # 8️⃣ Classificação de risco
    risco = "Baixo" if indice < 60 else "Moderado" if indice < 75 else "Alto"

    # 9️⃣ Prompt final para parecer estratégico
    prompt_final = f"""
    Gere parecer estratégico completo considerando:
    - Score aderência: {score_aderencia}
    - Horas semanais: {horas}
    - Status carga: {status_carga}
    - Score DISC: {disc_score}
    - Dificuldades: {dificuldades_classificadas}
    - Índice geral do cargo: {indice}
    - Classificação de risco: {risco}
    
    Inclua:
    - Diagnóstico estrutural
    - Análise de desvios
    - Avaliação comportamental
    - Riscos organizacionais
    - Recomendação detalhada de redistribuição
    - Atividades corretas para o cargo com tempo e frequência ideais
    - Conclusão executiva
    """

    # 10️⃣ Obter parecer do GPT
    parecer = ""
    try:
        if client:
            resposta = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role":"user","content":prompt_final}],
                temperature=0.3
            )
            parecer = resposta.choices[0].message.content
        else:
            parecer = "GPT não disponível. Retorno padrão: análise resumida."
    except:
        parecer = "Erro ao gerar parecer com GPT."

    # 11️⃣ Indicadores
    indicadores = {
        "score_aderencia": score_aderencia,
        "horas_semanais": horas,
        "status_carga": status_carga,
        "score_disc": disc_score,
        "indice_geral": indice,
        "risco": risco
    }

    return parecer, indicadores

# ============================================================
# GERAR PDF DO PARECER
# ============================================================

def gerar_pdf(parecer, nome):
    """
    Recebe:
    - parecer (texto)
    - nome do colaborador
    Cria arquivo PDF pronto para download
    """
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    nome_arquivo = f"{nome}_parecer.pdf"
    doc = SimpleDocTemplate(nome_arquivo)
    elements = []
    styles = getSampleStyleSheet()

    # Título
    elements.append(Paragraph("PARECER ESTRATÉGICO ORGANIZACIONAL", styles["Title"]))
    elements.append(Spacer(1, 0.5*inch))

    # Conteúdo linha a linha
    for linha in parecer.split("\n"):
        if linha.strip():  # evita parágrafos vazios
            elements.append(Paragraph(linha, styles["Normal"]))
            elements.append(Spacer(1, 0.2*inch))

    doc.build(elements)
    return nome_arquivo

# ============================================================
# PASTA BASE PARA FORMULÁRIOS (JSON)
# ============================================================
# Usamos 'dados_dir' para manter o padrão que já criamos
json_master = os.path.join(dados_dir, "formularios.json")

# Inicializa arquivo JSON se não existir
if not os.path.exists(json_master):
    with open(json_master, "w", encoding="utf-8") as f:
        json.dump([], f, ensure_ascii=False, indent=4)


import streamlit as st
import json
from datetime import datetime
from github import Github

# =========================================================
# 1. CONFIGURAÇÕES DE ACESSO (VIA STREAMLIT SECRETS)
# =========================================================
try:
    OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
  
    DB_TOKEN       = st.secrets["DB_TOKEN"]
    
    # Definimos o repositório direto aqui para evitar erro de Secret faltante
    REPO_NOME = "lucianohcl/formulario-colaborador"
    
except Exception as e:
    st.error(f"❌ Erro nos Secrets: A chave {e} não foi encontrada no painel do Streamlit.")
    st.stop()

# =========================================================
# 2. COLOQUE A FUNÇÃO AQUI (DEFINIÇÃO)
# =========================================================
def atualizar_rascunhos_do_github():
    # Inicializa rascunhos como dicionário vazio se não existir
    if "rascunhos" not in st.session_state:
        st.session_state["rascunhos"] = {}

    try:
        g = Github(DB_TOKEN)
        repo = g.get_repo(REPO_NOME)
        
        # Tenta acessar a pasta. Se falhar, tenta a raiz ""
        caminho = "rascunhos"
        try:
            contents = repo.get_contents(caminho)
        except:
            contents = repo.get_contents("")

        rascunhos_localizados = {}
        for content_file in contents:
            if content_file.name.endswith(".json"):
                try:
                    file_data = content_file.decoded_content.decode("utf-8")
                    dados_json = json.loads(file_data)
                    
                    # Normalização da Chave (Upper Case)
                    nome_raw = dados_json.get("colaborador") or dados_json.get("nome")
                    if nome_raw:
                        nome_chave = str(nome_raw).strip().upper() 
                        rascunhos_localizados[nome_chave] = dados_json
                except:
                    continue # Se um arquivo estiver com erro, não quebra os outros
        
        # IMPORTANTE: Só atualiza se encontrar algo, para não limpar o que já tem
        if rascunhos_localizados:
            st.session_state["rascunhos"] = rascunhos_localizados
            return True
        return False

    except Exception as e:
        # Se der erro de rede, mantém o que já estava na memória para não sumir tudo
        return False

# CHAME A FUNÇÃO AUTOMATICAMENTE NA INICIALIZAÇÃO
# Adicionamos uma trava para não ficar rodando toda hora sem necessidade
if "rascunhos" not in st.session_state or not st.session_state["rascunhos"]:
    atualizar_rascunhos_do_github()


# =========================================================
# 3. CHAME A EXECUÇÃO AQUI (INICIALIZAÇÃO)
# =========================================================
if "rascunhos" not in st.session_state:
    atualizar_rascunhos_do_github()

# 4. PARTE VISUAL (Daqui para baixo segue o seu st.title, etc)


# =========================================================
# 2. FUNÇÃO PARA SALVAR DADOS NO GITHUB
# =========================================================
def salvar_no_github(conteudo_dict, nome_arquivo):
    try:
        g = Github(DB_TOKEN)
        repo = g.get_repo(REPO_NOME)
        caminho_git = f"dados/{nome_arquivo}"
        
        json_string = json.dumps(conteudo_dict, ensure_ascii=False, indent=4)
        
        try:
            contents = repo.get_contents(caminho_git)
            repo.update_file(contents.path, f"Update: {nome_arquivo}", json_string, contents.sha)
        except:
            repo.create_file(caminho_git, f"Novo envio: {nome_arquivo}", json_string)
        
        return True
    except Exception as e:
        st.error(f"❌ Erro ao conectar com o GitHub: {e}")
        return False

# =========================================================
# 3. INTERFACE E LÓGICA DO FORMULÁRIO
# =========================================================


from github import Github
import json
import streamlit as st

def salvar_no_github(conteudo_dict, nome_arquivo):
    try:
        g = Github(st.secrets["DB_TOKEN"])
        repo = g.get_repo("lucianohcl/formulario-colaborador")
        caminho_git = f"dados/{nome_arquivo}"
        
        json_string = json.dumps(conteudo_dict, ensure_ascii=False, indent=4)
        
        try:
            contents = repo.get_contents(caminho_git)
            repo.update_file(contents.path, f"Update: {nome_arquivo}", json_string, contents.sha)
        except:
            repo.create_file(caminho_git, f"Novo envio: {nome_arquivo}", json_string)
        
        return True
    except Exception as e:
        st.error(f"❌ Erro ao conectar com o GitHub: {e}")
        return False







    # ============================================================
    # GARANTIA DE PERSISTÊNCIA (CARGA DOS DADOS)
    # ============================================================

    # Recarregamos os dados diretamente do disco/nuvem para garantir persistência total
    st.session_state["formularios"] = carregar_todos_formularios()

import streamlit as st
import pandas as pd
import plotly.express as px

# ============================================================
# 1. FUNÇÕES DE APOIO (CÁLCULOS E TRADUÇÃO)
# ============================================================

MAPA_DISC = {
    "A": "D", "B": "I", "C": "S", "D": "C"
}

def calcular_disc(respostas_disc):
    """Traduz A, B, C, D para D, I, S, C e calcula percentuais."""
    contagem = {"D": 0, "I": 0, "S": 0, "C": 0}
    for r in respostas_disc.values():
        perfil = MAPA_DISC.get(r)
        if perfil in contagem:
            contagem[perfil] += 1
    
    total = sum(contagem.values())
    if total > 0:
        percentuais = {k: round(v/total*100, 1) for k, v in contagem.items()}
        dominante = max(percentuais, key=percentuais.get)
    else:
        percentuais = {"D": 0, "I": 0, "S": 0, "C": 0}
        dominante = None
    return percentuais, dominante

# ============================================================
# 2. CARREGAMENTO E PERSISTÊNCIA
# ============================================================

# Garante que os dados estejam carregados na sessão
formularios = carregar_todos_formularios()
st.session_state["formularios"] = formularios

# ============================================================
# 3. PANORAMA COLETIVO (DENTRO DO EXPANDER)
# ============================================================

# ✅ Executa SÓ se o usuário clicou no menu "Perfil DISC"
if st.session_state.get("pagina") == "disc":  

    if formularios:
        # O 'expanded=False' garante que ele comece FECHADO
        with st.expander("📊 Ver Panorama Coletivo da Equipe", expanded=False):
            st.markdown("## 👥 Gestão Coletiva: Panorama da Equipe")
            
            lista_resultados = []
            atividades_coletivas = []

            # Processamento de todos os formulários carregados
            for f in formularios:
                res_percentual, _ = calcular_disc(f.get("disc", {}))
                lista_resultados.append(res_percentual)

                for a in f.get("atividades", []):
                    desc = a.get("Atividade Descrita", "").strip()
                    if desc:
                        atividades_coletivas.append(desc)

            if lista_resultados:
                # Criando DataFrame com a média de todos os perfis
                df_equipe = pd.DataFrame(lista_resultados).apply(pd.to_numeric, errors='coerce')
                medias = df_equipe.mean()
                
                # VARIÁVEIS DO GRUPO (A média real)
                dominante_grupo = medias.idxmax()
                menor_grupo = medias.idxmin()

                # --- Layout de Colunas ---
                col_txt, col_grf = st.columns([1, 1.5])
                
                with col_txt:
                    st.write("### 🧠 Insight do Grupo")
                    explicacoes = {
                        "D": "🔥 **Dominância:** Foco em metas e execução rápida.",
                        "I": "☀️ **Influência:** Comunicação e criatividade em alta.",
                        "S": "🌱 **Estabilidade:** Time leal, processual e resiliente.",
                        "C": "💎 **Conformidade:** Alta precisão técnica e perfeccionismo."
                    }
                    
                    st.info(f"**Perfil Dominante do Time:** {dominante_grupo}\n\n{explicacoes.get(dominante_grupo)}")
                    st.warning(f"**Menor Presença no Time:** {menor_grupo}")
                    st.caption(f"Análise baseada em {len(formularios)} formulários sincronizados.")

                with col_grf:
                    # Gráfico baseado nos dados agrupados
                    dados_plot = medias.reset_index()
                    dados_plot.columns = ["Tipo", "Media"]
                    
                    fig_eq = px.bar(
                        dados_plot, x="Tipo", y="Media", color="Tipo",
                        text_auto='.1f',
                        color_discrete_map={"D":"#FF4136", "I":"#FF851B", "S":"#2ECC40", "C":"#0074D9"}
                    )
                    fig_eq.update_layout(
                        template="plotly_white", height=280, showlegend=False,
                        yaxis_range=[0, 100], margin=dict(l=10, r=10, t=10, b=10)
                    )
                    st.plotly_chart(fig_eq, use_container_width=True)

                # --- Dificuldades de Adaptação ---
                st.divider()
                st.markdown(f"#### ⚠ Principais desafios de adaptação para o perfil {dominante_grupo}")
                
                # Lógica de ranking: Atividades menos compatíveis com o dominante do grupo
                compatibilidade_ativ = {
                    "D": ["decisão","meta","resultado","liderar","estratégia"],
                    "I": ["apresentar","comunicar","clientes","reunião"],
                    "S": ["suporte","atender","organizar","rotina","apoio"],
                    "C": ["analisar","dados","relatório","planilha","controle"]
                }

                ranking = []
                for ativ in atividades_coletivas:
                    texto = ativ.lower()
                    score = sum(p in texto for p in compatibilidade_ativ.get(dominante_grupo, []))
                    ranking.append((score, ativ))
                
                # Ordena pelo menor score (maior necessidade de adaptação)
                ranking.sort(key=lambda x: x[0])
                
                if ranking:
                    for _, atividade in ranking[:3]:
                        st.write(f"• {atividade}")
                else:
                    st.write("Nenhuma atividade descrita para análise.")

    else:
        st.info("Carregue formulários para habilitar o Panorama Coletivo.")



import streamlit as st
import pandas as pd
import json
from datetime import datetime
from github import Github

# =========================================================

# 1. CONFIGURAÇÕES E CONEXÃO

# =========================================================

st.set_page_config(page_title="Formulário Analítico", layout="wide")



# Puxa o Token dos segredos

TOKEN = st.secrets["DB_TOKEN"]



try:

    DB_USERNAME = st.secrets["DB_USERNAME"]

except Exception:

    DB_USERNAME = "lucianohcl"



# Conecta ao Github

g = Github(TOKEN)



# ESSE É O PONTO: Coloque o texto direto aqui para não dar erro de "REPO_NOME not defined"

repo = g.get_repo("lucianohcl/formulario-colaborador") 



# Se precisar do username para outra coisa, use direto dos secrets ou fixo:



if "rascunho" not in st.session_state: st.session_state["rascunho"] = {}

if "logado" not in st.session_state: st.session_state["logado"] = False



def val(chave, default=""):

    # Só tenta ler se o rascunho existir, senão ignora

    if "rascunho" in st.session_state:

        d = st.session_state["rascunho"]

        return d.get("campos", {}).get(chave, d.get(chave, default))

    return default


# =========================================================
# 2. IDENTIFICAÇÃO E CARREGAMENTO (VERSÃO BLINDADA)
# =========================================================
st.subheader("📋 Acesso ao Rascunho")
nome_input = st.text_input("DIGITE SEU NOME COMPLETO:").strip().upper()
# ADICIONE ESTA LINHA ABAIXO:
nome_digitado = nome_input

if not nome_input:
    st.info("Digite seu nome para começar.")
    st.stop()

nome_arq = f"rascunhos/{nome_input.replace(' ', '_')}.json"

if st.session_state.get("usuario_atual") != nome_input:
    st.session_state["usuario_atual"] = nome_input
    st.session_state["logado"] = False

confirmar = st.checkbox("✅ CLIQUE PARA CARREGAR MEUS DADOS")

if confirmar and not st.session_state.get("logado"):
    try:
        conteudo = repo.get_contents(nome_arq)
        dados_carregados = json.loads(conteudo.decoded_content.decode())
        # Garante que o rascunho seja EXATAMENTE o que está no GitHub
        st.session_state["rascunho"] = dados_carregados
        st.success("Dados recuperados!")
    except:
        st.session_state["rascunho"] = {"colaborador": nome_input, "campos": {}, "tabelas": {}, "disc": {}}
        st.info("Iniciando novo rascunho.")
    
    st.session_state["logado"] = True
    st.rerun()

# Se não confirmou, para aqui
if not st.session_state["logado"]:
    st.stop()

# =========================================================
# 3. FORMULÁRIO: CAMPOS DE TEXTO
# =========================================================
st.markdown("---")
col1, col2 = st.columns(2)
with col1:
    cargo = st.text_input("Cargo:", value=val("cargo"))
    depto = st.text_input("Departamento:", value=val("dep"))
    setor = st.text_input("Setor:", value=val("setor"))
with col2:
    chefe = st.text_input("Chefe imediato:", value=val("chefe"))
    unidade = st.text_input("Empresa / Unidade:", value=val("unidade"))
    escolaridade = st.text_input("Escolaridade:", value=val("escolaridade"))
    devolver_em = st.text_input("Devolver em:", value=val("devolver_em"))

cursos = st.text_area("Cursos Obrigatórios e Diferenciais:", value=val("cursos"))
objetivo = st.text_area("Em que consiste seu trabalho e qual seu Principal Objetivo:", value=val("objetivo"))


# =========================================================
# 4. MOTOR DE TABELAS (VERSÃO ULTRA-BLINDADA V3)
# =========================================================
def criar_editor(titulo, chave, col_p, col_e=None, nome_e=None):
    # --- 1. RESET DE MEMÓRIA (FORÇA O STREAMLIT A REDESENHAR O LAYOUT) ---
    # Se a v3 ainda não existe no estado da sessão, limpamos as versões antigas
    if f"ed_{chave}_v3" not in st.session_state:
        for k in list(st.session_state.keys()):
            if f"ed_{chave}" in k:
                del st.session_state[k]

    st.write(f"**{titulo}**")
    
    # 2. Puxa os dados do rascunho
    dados = st.session_state["rascunho"].get("tabelas", {}).get(chave, [])
    df = pd.DataFrame(dados)
    
    # 3. Define a Ordem Rígida (O Setor/Impacto TEM que ser a segunda coluna)
    if col_e:
        cols_finais = [col_p, col_e, "Horas", "Minutos", "Frequência"]
    else:
        cols_finais = [col_p, "Horas", "Minutos", "Frequência"]
    
    # 4. Limpeza e Reindex (Remove colunas fantasmas e organiza a ordem)
    df = df.fillna("").astype(str)
    for c in df.columns:
        df[c] = df[c].str.strip()
    
    # O reindex com columns=cols_finais descarta qualquer coluna que não esteja na lista
    df = df.reindex(columns=cols_finais, fill_value="")
    
    # 5. Garante as 15 linhas fixas
    if len(df) < 15:
        faltam = 15 - len(df)
        extras = pd.DataFrame([{c: "" for c in cols_finais} for _ in range(faltam)])
        df = pd.concat([df, extras], ignore_index=True)
    
    # Trava em 15 linhas e na ordem correta
    df = df[cols_finais].head(15)

    # 6. Configuração Visual dos Seletores
    cfg = {
        col_p: st.column_config.TextColumn("Descrição", width="large"),
        "Frequência": st.column_config.SelectboxColumn(
            options=["", "DVD", "D", "S", "Q", "M", "T", "A"], 
            width="small"
        ),
        "Horas": st.column_config.SelectboxColumn(
            options=[""] + [f"{i} h" for i in range(25)], 
            width="small"
        ),
        "Minutos": st.column_config.SelectboxColumn(
            options=[""] + [f"{i} min" for i in range(0, 60, 5)], 
            width="small"
        ),
    }
    if col_e: 
        cfg[col_e] = st.column_config.TextColumn(nome_e, width="medium")
        
    # 7. Renderização com a Key v3 (O segredo para resetar o visual)
    return st.data_editor(
        df, 
        key=f"ed_{chave}_v3", 
        column_config=cfg, 
        use_container_width=True,
        num_rows="fixed"
    )

# Chamadas das tabelas
e_alta = criar_editor("🚀 Alta Complexidade", "alta", "Atividade")
e_normal = criar_editor("📋 Complexidade Normal", "normal", "Atividade")
e_baixa = criar_editor("⏳ Baixa Complexidade", "baixa", "Atividade")
e_dif = criar_editor("⚠️ Dificuldades", "dificuldades", "Dificuldade", "Setor Envolvido")
e_sug = criar_editor("💡 Sugestões", "sugestoes", "Sugestão", "Impacto Esperado")


# =========================================================
# 6. PERFIL DISC (PERSISTÊNCIA GARANTIDA - CORRIGIDO)
# =========================================================
st.markdown("---")
st.subheader("📊 Questionário DISC")

# =========================================================
# 🔥 GARANTE ESTRUTURA
# =========================================================
if "rascunho_atual" not in st.session_state:
    st.session_state["rascunho_atual"] = {}

if "disc" not in st.session_state["rascunho_atual"]:
    st.session_state["rascunho_atual"]["disc"] = {}

# Recupera o dicionário salvo do rascunho (Garante chaves como string)
disc_data = {
    str(k): v for k, v in st.session_state["rascunho_atual"]["disc"].items()
}

# Pegamos o nome do colaborador para resetar os campos se mudar de pessoa
nome_colab = st.session_state.get("nome_colaborador", "novo")

# =========================================================
# 📋 LISTA DE PERGUNTAS (EXATA PARA ESPELHAMENTO)
# =========================================================
perguntas_disc = [
    "Quando surge um problema inesperado: (A) Age rápido | (B) Comunica a todos | (C) Analisa riscos | (D) Segue processo",
    "Em situações de pressão: (A) Foca no resultado | (B) Mantém o otimismo | (C) Mantém a calma | (D) Busca precisão",
    "Ao receber tarefa difícil: (A) Aceita o desafio | (B) Busca ajuda social | (C) Planeja passos | (D) Estuda as regras",
    "No trabalho em equipe: (A) Lidera o grupo | (B) Motiva os colegas | (C) Apoia os outros | (D) Organiza as tarefas",
    "Em reuniões: (A) Vai direto ao ponto | (B) Interage e brinca | (C) Escuta mais | (D) Anota detalhes",
    "Ao lidar com conflitos: (A) Enfrenta direto | (B) Tenta apaziguar | (C) Evita o confronto | (D) Usa lógica e fatos",
    "Seu ritmo de trabalho: (A) Rápido/Impaciente | (B) Rápido/Entusiasmado | (C) Calmo/Constante | (D) Metódico/Cauteloso",
    "Prefere tarefas: (A) Desafiadoras | (B) Variadas e sociais | (C) Rotineiras e seguras | (D) Técnicas e detalhadas",
    "Seu foco principal: (A) Resultados | (B) Relacionamentos | (C) Estabilidade | (D) Qualidade e Processos",
    "Ao decidir, você é: (A) Decidido e firme | (B) Impulsivo e intuitivo | (C) Cuidadoso e lento | (D) Lógico e analítico",
    "Confia mais em: (A) Sua intuição | (B) Opinião alheia | (C) Experiência passada | (D) Dados e provas",
    "Prefere decisões: (A) Independentes | (B) Em grupo | (C) Consensuais | (D) Baseadas em normas",
    "Estilo de organização: (A) Prático | (B) Criativo/Bagunçado | (C) Tradicional | (D) Muito organizado",
    "Lida melhor com: (A) Mudanças rápidas | (B) Novas ideias | (C) Rotinas claras | (D) Regras rígidas",
    "Prefere trabalhar: (A) Sozinho/Comando | (B) Ambiente festivo | (C) Ambiente tranquilo | (D) Ambiente silencioso",
    "Seu ponto forte: (A) Coragem | (B) Comunicação | (C) Paciência | (D) Organização",
    "Você se considera: (A) Dominante | (B) Influente | (C) Estável | (D) Conforme/Analítico",
    "Se motiva por: (A) Poder/Bônus | (B) Reconhecimento | (C) Segurança/Paz | (D) Conhecimento Técnico",
    "Reação a cobranças: (A) Mais esforço | (B) Desculpas criativas | (C) Ansiedade | (D) Argumentos técnicos",
    "Ambiente ideal: (A) Competitivo | (B) Amigável | (C) Previsível | (D) Disciplinado",
    "Ao lidar com feedback: (A) Aceita e ajusta | (B) Comenta e debate | (C) Analisa e planeja | (D) Segue regras",
    "Como prefere aprender: (A) Fazendo | (B) Interagindo | (C) Observando | (D) Estudando materiais",
    "Gestão de tempo: (A) Prioriza resultados | (B) Mantém relações | (C) Planeja com cuidado | (D) Segue processos",
    "Como se comunica: (A) Direto e objetivo | (B) Amigável e motivador | (C) Calmo e ponderado | (D) Técnico e detalhista"
]

# =========================================================
# 🎯 DEFINIÇÕES E RENDER (SINCRONIZADO)
# =========================================================
opcoes = ["A", "B", "C", "D"]
respostas_disc_final = {} # Nome correto para evitar o NameError

for i, pergunta in enumerate(perguntas_disc):
    chave = str(i)
    
    # Busca a letra no banco. Se não houver, retorna None para ficar desmarcado
    letra_banco = disc_data.get(chave)
    
    # Define o índice da bolinha (A=0, B=1, C=2, D=3)
    idx = opcoes.index(letra_banco) if letra_banco in opcoes else None

    respostas_disc_final[chave] = st.radio(
        f"**{i+1}. {pergunta}**",
        options=opcoes,
        index=idx,
        horizontal=True,
        key=f"disc_{nome_colab}_{i}" # Key dinâmica para resetar ao trocar colaborador
    )

# =========================================================
# 💾 PERSISTÊNCIA AUTOMÁTICA
# =========================================================
st.session_state["rascunho_atual"]["disc"] = respostas_disc_final



# =========================================================
# 6. SALVAMENTO (GITHUB)
# =========================================================
if st.button("💾 SALVAR TUDO", use_container_width=True):

    # --- VALIDAÇÃO CRÍTICA ---
    if not nome_input or len(nome_input) < 3:
        st.error("⚠️ Erro: Nome do colaborador está vazio ou inválido.")
        st.stop()

    # 1. Monta o payload (O "corpo" do arquivo)
    payload = {
        "colaborador": nome_input, 
        "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        "campos": {
            "cargo": cargo, "dep": depto, "setor": setor, 
            "chefe": chefe, "unidade": unidade, 
            "escolaridade": escolaridade, "devolver_em": devolver_em, 
            "cursos": cursos, "objetivo": objetivo
        },
        "tabelas": {
            "alta": e_alta.to_dict("records"), 
            "normal": e_normal.to_dict("records"), 
            "baixa": e_baixa.to_dict("records"), 
            "dificuldades": e_dif.to_dict("records"), 
            "sugestoes": e_sug.to_dict("records")
        },
        "disc": respostas_disc_final  # <--- AGORA O NOME ESTÁ IGUAL AO QUE VOCÊ CRIOU
    }

    # --- CONFIGURAÇÃO DO NOME DO ARQUIVO (A CHAVE DO SUCESSO) ---
    nome_limpo = nome_input.strip().replace(" ", "_").upper()
    caminho_github = f"rascunhos/{nome_limpo}.json" # Garante que vai para a pasta rascunhos

    try:
        # --- AÇÃO 1: GITHUB (ENVIO REAL PARA O REPOSITÓRIO) ---
        conteudo_json = json.dumps(payload, indent=4, ensure_ascii=False)
        
        try:
            # Tenta atualizar se já existir
            f = repo.get_contents(caminho_github)
            repo.update_file(f.path, f"Update: {nome_input}", conteudo_json, f.sha)
        except:
            # Cria novo se não existir
            repo.create_file(caminho_github, f"Novo: {nome_input}", conteudo_json)
        
        # Salva no estado da sessão para uso imediato
        st.session_state["rascunho"] = payload
        st.success(f"✅ {nome_input} salvo com sucesso no GitHub!")

        # === COLOQUE AQUI O BOTÃO DE DOWNLOAD ===
        st.download_button(
            label="📥 Baixar Arquivo JSON do Rascunho",
            data=conteudo_json.encode('utf-8'), # Usa a variável que você já criou acima
            file_name=f"{nome_limpo}.json",
            mime="application/json",
            use_container_width=True
        )
        # =========================================

        # --- AÇÃO 2: SHEETS ---
        if enviar_para_sheets(payload):
            st.toast("📊 Espelhado no Sheets!", icon="📈")
        else:
            st.warning("⚠️ Salvo no GitHub, mas Sheets não respondeu.")

    except Exception as e:
        st.error(f"⚠️ Erro ao salvar: {e}")
        # Botão de emergência caso o GitHub falhe
        st.download_button(
            label="📥 Baixar Backup de Emergência",
            data=json.dumps(payload, indent=4, ensure_ascii=False).encode('utf-8'),
            file_name=f"{nome_limpo}_EMERGENCIA.json",
            mime="application/json",
            use_container_width=True
        )
# Versao_Final_06_04 
